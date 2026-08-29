# -*- coding: utf-8 -*-
"""Insert current two-room figures into the AU thesis and refresh TOC."""
from __future__ import annotations

import shutil
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(r"c:\Users\User\OneDrive\Documents\ULTRA_Project")
DOC = ROOT / "Thesis_Paper" / "Air_Thesis_Formate" / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"
FIGS = ROOT / "validate" / "dual_index_routing" / "figures"
FIG1 = FIGS / "fig_two_rooms_lights.png"
FIG_LAYERS = FIGS / "fig_three_evaluation_layers.png"
FIG_CUE = FIGS / "fig_cue_split.png"
FIG_P5 = FIGS / "fig_heldout_p5.png"


def insert_after(paragraph, text: str, style_name: str):
    el = deepcopy(paragraph._p)
    paragraph._p.addnext(el)
    np = Paragraph(el, paragraph._parent)
    try:
        np.style = style_name
    except Exception:
        pass
    if np.runs:
        np.runs[0].text = text
        for r in np.runs[1:]:
            r.text = ""
    else:
        np.add_run(text)
    return np


def add_picture_after(paragraph, image: Path, caption: str):
    cap = insert_after(paragraph, caption, "Caption")
    try:
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
    pic = insert_after(paragraph, "", "Normal")
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if pic.runs:
        pic.runs[0].text = ""
    run = pic.add_run()
    run.add_picture(str(image), width=Inches(5.9))
    return pic, cap


def replace_zip_image(docx_path: Path, member: str, png: Path):
    tmp = docx_path.with_name(docx_path.stem + "._img.tmp.docx")
    names = []
    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        if member not in names:
            raise FileNotFoundError(member)
        with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for name in names:
                data = png.read_bytes() if name == member else zin.read(name)
                zout.writestr(name, data)
    shutil.copy2(tmp, docx_path)
    tmp.unlink(missing_ok=True)


def main():
    if not DOC.exists():
        raise SystemExit(f"missing {DOC}")
    for p in (FIG1, FIG_LAYERS, FIG_CUE, FIG_P5):
        if not p.exists():
            raise SystemExit(f"missing {p}")

    replace_zip_image(DOC, "word/media/image1.png", FIG1)
    print("replaced Figure 1 image (word/media/image1.png)")

    doc = Document(str(DOC))
    anchors = {}
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("5.13 Frozen Phase 3B") and "anchors" not in str(anchors.get("513")):
            anchors["513"] = p
        if t.startswith("5.16 Frozen held-out trap classification"):
            anchors["516"] = p
        if t.startswith("5.17 Frozen held-out dual-index P@5"):
            anchors["517"] = p

    if "513" in anchors:
        add_picture_after(
            anchors["513"],
            FIG_LAYERS,
            "Figure 18. Three evaluation layers that must not be mixed: development/CV, frozen Phase 3B (86/84), and frozen traps (60/20/50).",
        )
        print("inserted Figure 18 after 5.13")
    if "516" in anchors:
        add_picture_after(
            anchors["516"],
            FIG_CUE,
            "Figure 19. Held-out trap cue split (n = 40): SVM 18/18 when why/how/fact words fire; 27.27% for both systems otherwise.",
        )
        print("inserted Figure 19 after 5.16")
    if "517" in anchors:
        add_picture_after(
            anchors["517"],
            FIG_P5,
            "Figure 20. Held-out dual-index graded P@5 (400 judgments): word count 36.50%, always-headline 35.00%, SVM 33.00%.",
        )
        print("inserted Figure 20 after 5.17")

    # TOC figure list lines
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("Figure 17. Consolidated robustness"):
            insert_after(p, "Figure 20. Held-out dual-index graded P@5", p.style.name)
            insert_after(p, "Figure 19. Held-out trap cue split", p.style.name)
            insert_after(p, "Figure 18. Three evaluation layers (do not mix)", p.style.name)
            break

    doc.save(str(DOC))
    print("saved", DOC)


if __name__ == "__main__":
    main()
