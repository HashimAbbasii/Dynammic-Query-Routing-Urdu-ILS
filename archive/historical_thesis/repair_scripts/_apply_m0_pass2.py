# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document

HERE = Path(__file__).resolve().parent
SRC = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def main():
    doc = Document(str(SRC))
    jobs = [
        (
            "At the retrieval-result level, Precision at rank k (P@k) — used here at k=15",
            "At the retrieval-result level this thesis uses two official IR metrics that must not be mixed: ExactSource Hit@5 (known-item recovery of a pre-assigned source document) and human Success@5 (at least one A or B in the Top-5). Development notebooks also reported P@15 for a small Layer A retrieval experiment; that P@15 is not official M0 performance. Mean Reciprocal Rank is reported for U as a secondary usefulness statistic.",
        ),
        (
            "Development retrieval quality is reported as P@15. Frozen dual-index quality is graded P@5",
            "Official M0 retrieval quality is ExactSource Hit@5 (n=78 and K) and human Success@5 (U). Historical Layer A also reported development P@15 and frozen dual-index graded P@5 (Relevant = 1, Partially relevant = 0.5, Not relevant = 0):",
        ),
        (
            "The diagnostic external validation on 50 unseen native-Urdu queries (development/robustness layer). The deployed V2 model records 98.00% on this 50-query diagnostic set (training_info.json); the frozen generalization test is Phase 3B (Section 5.13).",
            "Diagnostic SVM routing check on 50 native-Urdu queries (development/robustness layer). V2 records 98.00% routing accuracy on this set. This is not official M0 ExactSource Hit@5. Official IR tests are Sections 5.19–5.22.",
        ),
    ]
    # also match caption without "The" prefix
    jobs.append(
        (
            "Diagnostic external validation on 50 unseen native-Urdu queries (development/robustness layer). The deployed V2 model records 98.00% on this 50-query diagnostic set (training_info.json); the frozen generalization test is Phase 3B (Section 5.13).",
            "Diagnostic SVM routing check on 50 native-Urdu queries (development layer). V2 98.00% is routing accuracy, not official M0 Hit@5. Official IR tests are Sections 5.19–5.22.",
        )
    )
    hit, miss = 0, []
    used = set()
    for prefix, new in jobs:
        found = False
        for i, p in enumerate(doc.paragraphs):
            if i in used:
                continue
            t = p.text or ""
            if prefix in t or t.startswith(prefix[:60]):
                set_text(p, new)
                used.add(i)
                hit += 1
                found = True
                break
        if not found:
            miss.append(prefix[:70])
    # closing remarks
    for p in doc.paragraphs:
        t = p.text or ""
        if t.startswith("This thesis set out to answer a narrow question") or t.startswith("We replaced ULTRA"):
            set_text(
                p,
                "This thesis set out to measure script-aware Urdu news retrieval under a freeze. M0 achieves 87.18% ExactSource Hit@5 on the development/validation known-item protocol, 67.50% on new known-item queries, and 57.50% human Success@5 on naturalistic queries. Roman Urdu remains the main limitation. That is enough for an honest MS contribution and not enough to claim universal effectiveness.",
            )
            hit += 1
            break
    doc.save(str(SRC))
    print("second_pass hit", hit, "miss", miss)


if __name__ == "__main__":
    main()
