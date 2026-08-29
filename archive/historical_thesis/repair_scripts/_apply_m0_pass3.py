# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document

SRC = Path(__file__).resolve().parent / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def main():
    doc = Document(str(SRC))
    jobs = [
        (
            "Prefer classical ML routing over LLM-based routing for well-scoped decisions.",
            "Do not treat 87.18% ExactSource Hit@5 as a user-facing SLA. For native-script Urdu news on this collection, frozen M0 is a reasonable lexical baseline. For Roman Urdu traffic, expect a large drop (U Success@5 = 6/18 in the sealed sample). If the system is changed, freeze it first and seal a new test; K, U, and H001–H040 are burned for tuning.",
        ),
        (
            "Compare a learned router against the strongest simple rule, not only against a known-broken rule.",
            "Keep ExactSource Hit@5 and human Success@5 separate. Do not average 87.18%, 67.50%, and 57.50%. Layer A SVM classification (86/84) is not official M0 retrieval performance.",
        ),
        (
            "Large-scale, community benchmark construction. A natural next step is to construct and publicly release a substantially larger, independently annotated Urdu and Roman Urdu query benchmark",
            "Do not tune BM25, the dictionary, routing, or Method D on U001–U040, K001–K040, or H001–H040. Future Roman work needs a new sealed test. A larger independently annotated Urdu and Roman Urdu query benchmark would also enable more robust comparison than n=40.",
        ),
        (
            "Need features beyond cue words, and a corpus that answers why.",
            "Priority directions that require a new sealed evaluation: better Roman query–document matching than Method D; mixed-script evaluation with more than four queries; a second annotator on a new naturalistic set.",
        ),
        (
            "The question was whether Urdu search should still flip rooms with 150 letters.",
            "This thesis measured script-aware Urdu news retrieval under a freeze. M0 achieves 87.18% ExactSource Hit@5 on the development/validation known-item protocol, 67.50% on new known-item queries, and 57.50% human Success@5 on naturalistic queries. Roman Urdu remains the main limitation. That is enough for an honest MS contribution and not enough to claim universal effectiveness.",
        ),
    ]
    hit, miss = 0, []
    used = set()
    for prefix, new in jobs:
        found = False
        for i, p in enumerate(doc.paragraphs):
            if i in used:
                continue
            t = p.text or ""
            if t.startswith(prefix) or prefix[:50] in t[:120]:
                set_text(p, new)
                used.add(i)
                hit += 1
                found = True
                break
        if not found:
            miss.append(prefix[:70])
    doc.save(str(SRC))
    print("pass3 hit", hit, "miss", miss)


if __name__ == "__main__":
    main()
