# -*- coding: utf-8 -*-
"""Replace typed TOC / lists with real Word fields so Update Table works."""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

HERE = Path(__file__).resolve().parent
SRC = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"
BAK = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.pre_tocfields.bak.docx"

CAPTION_RE = re.compile(r"^(Figure|Table)\s+[\d.]+\.\s*(.*)$", re.S)


def append_field(paragraph: Paragraph, instruction: str, placeholder: str) -> None:
    def add_child(child) -> None:
        run = OxmlElement("w:r")
        run.append(child)
        paragraph._p.append(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    add_child(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction.strip()} "
    add_child(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    add_child(sep)

    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = placeholder
    add_child(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    add_child(end)


def clear_runs(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag == qn("w:r"):
            paragraph._p.remove(child)


def delete_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def find_para(doc: Document, title: str) -> Paragraph:
    for p in doc.paragraphs:
        if (p.text or "").strip() == title:
            return p
    raise SystemExit(f"missing heading: {title}")


def delete_until(doc: Document, start_title: str, end_title: str) -> Paragraph:
    paras = list(doc.paragraphs)
    start_i = end_i = None
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if t == start_title and start_i is None:
            start_i = i
        elif start_i is not None and t == end_title:
            end_i = i
            break
    if start_i is None or end_i is None:
        raise SystemExit(f"could not find block {start_title!r} .. {end_title!r}")
    for p in paras[start_i + 1 : end_i]:
        delete_paragraph(p)
    return paras[start_i]


def insert_empty_after(paragraph: Paragraph) -> Paragraph:
    new_el = deepcopy(paragraph._p)
    paragraph._p.addnext(new_el)
    np = Paragraph(new_el, paragraph._parent)
    try:
        np.style = "Normal"
    except Exception:
        pass
    clear_runs(np)
    return np


def add_seq_captions(doc: Document) -> tuple[int, int]:
    n_fig = n_tab = 0
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if style != "Caption":
            continue
        m = CAPTION_RE.match((p.text or "").strip())
        if not m:
            continue
        kind, title = m.group(1), m.group(2).strip()
        clear_runs(p)
        p.add_run(f"{kind} ")
        append_field(p, f"SEQ {kind} \\* ARABIC", "1" if kind == "Figure" else "1")
        p.add_run(f". {title}")
        if kind == "Figure":
            n_fig += 1
        else:
            n_tab += 1
    return n_fig, n_tab


def enable_update_on_open(doc: Document) -> None:
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)
    else:
        existing.set(qn("w:val"), "true")


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"missing {SRC}")
    shutil.copy2(SRC, BAK)

    doc = Document(str(SRC))
    n_fig, n_tab = add_seq_captions(doc)
    print("seq captions", n_fig, "figures", n_tab, "tables")

    toc_h = delete_until(doc, "Table of Contents", "List of Figures")
    toc_p = insert_empty_after(toc_h)
    append_field(
        toc_p,
        'TOC \\o "1-3" \\h \\z \\u',
        "Open this file in desktop Word and click Yes when asked to update fields.",
    )

    fig_h = delete_until(doc, "List of Figures", "List of Tables")
    fig_p = insert_empty_after(fig_h)
    append_field(
        fig_p,
        'TOC \\c "Figure" \\h \\z',
        "Open this file in desktop Word and click Yes when asked to update fields.",
    )

    tab_h = delete_until(doc, "List of Tables", "Symbols and Abbreviations")
    tab_p = insert_empty_after(tab_h)
    append_field(
        tab_p,
        'TOC \\c "Table" \\h \\z',
        "Open this file in desktop Word and click Yes when asked to update fields.",
    )

    enable_update_on_open(doc)
    out = SRC
    try:
        doc.save(str(out))
    except PermissionError:
        out = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA_fields.docx"
        doc.save(str(out))
        print("original is open/locked; wrote", out)
    else:
        print("saved", out)
    print("backup", BAK)


if __name__ == "__main__":
    main()
