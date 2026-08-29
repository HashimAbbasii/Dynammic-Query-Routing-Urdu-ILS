# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document

HERE = Path(__file__).resolve().parent
d = Document(str(HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"))
keys = (
    "5.17 Frozen",
    "5.18 Official",
    "5.19 Phase",
    "5.20 Phase",
    "5.21 Phase",
    "5.22 Phase",
    "5.23 Final",
    "5.24 Discussion",
    "Chapter 6:",
    "4.13 Official",
    "Table of Contents",
    "List of Figures",
    "List of Tables",
    "Do not retune M0",
    "Table 22.",
    "Table 3. Phase 2 n=78",
    "87.18",
    "67.50",
    "57.50",
)
out = []
out.append("paragraphs=%d tables=%d" % (len(d.paragraphs), len(d.tables)))
for i, p in enumerate(d.paragraphs):
    t = p.text or ""
    if not any(k in t for k in keys):
        continue
    s = (p.style.name if p.style else "")[:28]
    snippet = t[:180].replace("\u2192", "->")
    out.append("[%d|%s] %s" % (i, s, snippet))

# Check order: 5.18 must be after body 5.17 heading and before body Ch6
idx_517 = idx_518 = idx_ch6 = None
for i, p in enumerate(d.paragraphs):
    t = (p.text or "").strip()
    st = p.style.name if p.style else ""
    if t.startswith("5.17 Frozen") and st.startswith("Heading"):
        idx_517 = i
    if t.startswith("5.18 Official") and st.startswith("Heading"):
        idx_518 = i
    if t == "Chapter 6: Conclusions and Recommendations" and st.startswith("Heading"):
        idx_ch6 = i
out.append("ORDER heading2 5.17=%s 5.18=%s body_ch6=%s" % (idx_517, idx_518, idx_ch6))
if idx_517 is not None and idx_518 is not None and idx_ch6 is not None:
    ok = idx_517 < idx_518 < idx_ch6
    out.append("ORDER_OK=%s" % ok)

# leftover TOC mash
mash = 0
for p in d.paragraphs:
    t = p.text or ""
    if "5.17 Frozen held-out dual-index P@5 (400 judgments)" in t and "5.18" in t:
        mash += 1
out.append("TOC_MASH_COUNT=%d" % mash)

# official numbers still present
blob = "\n".join((p.text or "") for p in d.paragraphs)
for phrase in (
    "68/78 = 87.18%",
    "27/40 = 67.50%",
    "23/40 = 57.50%",
    "M1 is a gate-passing candidate",
    "111,860",
):
    out.append("HAS %s: %s" % (phrase, phrase in blob))

Path(HERE / "_audit_verify.txt").write_text("\n".join(out), encoding="utf-8")
print("wrote _audit_verify.txt lines", len(out))
