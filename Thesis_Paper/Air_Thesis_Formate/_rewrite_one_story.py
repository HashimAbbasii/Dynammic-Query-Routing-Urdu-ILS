# -*- coding: utf-8 -*-
"""Align the AU thesis to one story: two rooms, lights, 86/84, 60/20, cue-split, P@5."""
from __future__ import annotations

import json
import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

HERE = Path(__file__).resolve().parent
SRC = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"
BAK = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.pre_onestory.bak.docx"
REPO = HERE.parents[1]
DICT_PATH = REPO / "models" / "roman_urdu_dict_expanded.json"


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def insert_after(paragraph, text: str, style_name: str):
    new_el = deepcopy(paragraph._p)
    paragraph._p.addnext(new_el)
    np = Paragraph(new_el, paragraph._parent)
    try:
        np.style = style_name
    except Exception:
        pass
    set_text(np, text)
    return np


def replace_prefixed(doc, jobs):
    """jobs: list of (prefix, new_text). First matching paragraph wins."""
    hit, miss = [], []
    used = set()
    for prefix, new in jobs:
        found = False
        for i, p in enumerate(doc.paragraphs):
            if i in used:
                continue
            t = p.text or ""
            if re.search(r"\t\d+\s*$", t) and len(t) < 140:
                continue
            if t.startswith(prefix) or (len(prefix) > 40 and prefix in t):
                set_text(p, new)
                used.add(i)
                hit.append(prefix[:60])
                found = True
                break
        if not found:
            miss.append(prefix[:80])
    return hit, miss


def dict_count() -> int:
    if not DICT_PATH.exists():
        return 198
    with DICT_PATH.open(encoding="utf-8") as f:
        return len(json.load(f))


def main():
    n_dict = dict_count()
    if not SRC.exists():
        raise SystemExit(f"missing thesis: {SRC}")
    if not BAK.exists():
        shutil.copy2(SRC, BAK)

    doc = Document(str(SRC))
    print("dict_count", n_dict)

    jobs = [
        # --- Abstract ---
        (
            "Urdu information retrieval (IR) systems face persistent challenges arising from script variability",
            "Urdu information retrieval (IR) systems face persistent challenges arising from script variability, morphological complexity, and the widespread use of Roman Urdu in user queries. The ULTRA framework still switches retrieval with a static character-length rule (θ = 150). That rule is the wrong switch for this query population: it almost never fires LONG, and it does not say whether the user needs a headline or the article. This thesis replaces that tape with a learned router that opens one of two retrieval rooms: a headline-only semantic index (SHORT: a headline is enough) or a full-article semantic index (LONG: the user needs the body). A Support Vector Machine (SVM) makes the room decision. Confidence acts as traffic lights: HIGH (≥ 85%) searches only the chosen room; MEDIUM (60–85%) mixes both rooms at equal weight; LOW (< 60%) expands the query and then mixes both rooms. Roman Urdu queries are mapped through an expanded transliteration dictionary (198 on-disk word pairs; earlier drafts cited 179) before routing and search. Three evaluation layers must not be mixed. Development and cross-validation show that the routing task is learnable, including 100% on some splits versus 50% for θ = 150; those figures are not the generalization claim. Frozen Phase 3B on the V2 eight-feature SVM (409 training queries) is 86.00% versus 84.00% for a six-word baseline (exact McNemar p = 1.0000). After 38 trap queries labelled by retrieval need were added, a twelve-feature SVM was trained (447 queries). On a new frozen 40-query trap set that was never used for training, that model scores 60.00% against 20.00% for word count and 50.00% for θ = 150 (McNemar 16–0, exact p < 0.001). The classification gain appears only when why/how/fact cue words fire (18/18); without those cues both systems score 27.27%. On the same 40 queries, 400 dual-index headline judgments yield graded P@5 of 36.50% (word count), 35.00% (always headlines / θ = 150), 34.25% (always full text), and 33.00% (SVM). Routing labels therefore beat length rules on traps; dual-index P@5 does not yet improve, because many why-queries still retrieve “price rose” headlines rather than causal articles. A lightweight SVM also matched LLM-style routers in a development comparison at sub-millisecond local cost. The frozen 86/84 and 60/20 tables, the cue split, and the P@5 table are the headline results.",
        ),
        # --- Ch1 intro ---
        (
            "This thesis extends the ULTRA framework by replacing its static, length-based routing threshold with an adaptive, dynamic query routing mechanism.",
            "This thesis extends ULTRA by replacing the θ = 150 character tape as a whole: the decision and the search that must follow it. SHORT opens a headline semantic room; LONG opens a full-article semantic room. An SVM learns that room choice. Lights mix or expand when the SVM is unsure. Roman Urdu is dictionary-transliterated first. Routing is treated as a learnable problem, and is evaluated both as a classification decision and as dual-index P@5.",
        ),
        (
            "Existing Urdu information retrieval systems, including the ULTRA framework on which this thesis builds, predominantly employ static, threshold-based routing",
            "Existing Urdu IR systems, including ULTRA, still treat routing as a length tape. θ = 150 cannot tell a short query that needs a story from a long query that needs one fact, and it cannot open a different index for those two needs. A six-word rule is a stronger simple alternative, but it is still a tape: it does not encode headline-enough versus need-the-article. The second failure is script: a Roman Urdu fragment and a native-Urdu question of the same length are processed as if they were the same object. The third failure is uncertainty: a static rule never says “I am unsure,” so it cannot mix rooms or expand the query. This thesis addresses all three by learning the room, lighting the mix, and transliterating Roman Urdu before search.",
        ),
        (
            "This problem manifests concretely in at least three ways.",
            "Those failures show up in three concrete ways. First, character length does not measure retrieval need: a three-word why-query may need the article, while a long fact query may need only a headline. Second, script type is invisible to θ = 150, so Roman Urdu never gets a dedicated path. Third, every query is routed with equal false certainty, which blocks HIGH / MEDIUM / LOW lights. Chapter 4 implements the two rooms and the lights; Chapter 5 reports frozen classification and frozen dual-index P@5 separately.",
        ),
        (
            "A search system that cannot reliably decide how to handle a query",
            "A search system that cannot decide which room to open will under-serve users even if the encoder is strong. For Urdu, a large share of informal queries are Roman rather than Perso-Arabic. A tape on character length cannot distinguish a short native-Urdu fact from a short Roman fragment, and it cannot send them to different indexes. The motivation is therefore twofold: show that a small feature set can beat a broken character rule and can beat a six-word tape on need-based trap labels; and show that this can run locally, without an LLM call per query. Dual-index P@5 is the test of whether the room choice actually helps retrieval. In this thesis it does not yet beat word count — that limit is part of the claim, not a result to hide.",
        ),
        (
            "Synthesizing the discussion above, the specific research gap this thesis addresses can be stated precisely:",
            "The gap is joint, not three separate papers. No Urdu IR system reviewed here (i) learns SHORT versus LONG as headline-enough versus need-the-article, (ii) opens two semantic rooms for that decision, and (iii) uses calibrated lights to mix or expand when the router is unsure, while also handling Roman Urdu inside the same pipeline. English query routing exists; Urdu resources exist; the intersection — learned dual-index routing with honest frozen tests — did not. Chapter 2 details that gap.",
        ),
        (
            "RQ1: Can query-level structural and linguistic features",
            "RQ1: Can a small set of query features (script composition, length, and later why/how/fact cues) predict which room a user needs — headlines only versus the full article — better than a character or word-count tape?",
        ),
        (
            "RQ2: Does a dynamic, machine learning-based query routing mechanism outperform static threshold-based routing",
            "RQ2: Does that learned router beat θ = 150 and a six-word rule on frozen labels, and does choosing the matching room raise dual-index P@5? Phase 3B answers the first part for V2 (86% vs 84%). A later frozen 40-query trap set answers it for need-based labels (60% vs 20%). Dual-index P@5 on the same 40 queries answers the retrieval part (Section 5.17).",
        ),
        (
            "RQ3: Can a confidence-based tiered routing strategy further improve system reliability",
            "RQ3: Can calibrated confidence act as traffic lights — HIGH: one room; MEDIUM: mix both rooms; LOW: expand the query then mix — without replacing the frozen accuracy tables?",
        ),
        (
            "Dynamic query classification: Design and train a dynamic SVM-based classifier that predicts optimal query routing decisions using eight structural and linguistic features",
            "Dynamic query classification: Train an SVM that labels SHORT (headline room) versus LONG (full-article room) for Urdu and Roman Urdu queries.",
        ),
        (
            "Roman Urdu support: Develop and expand a Roman Urdu transliteration dictionary (from 30 to 179 words)",
            f"Roman Urdu support: Detect Latin-script Urdu and transliterate with an expanded dictionary ({n_dict} on-disk pairs) before routing and search.",
        ),
        (
            "Confidence-based tiered routing: Implement a three-tier confidence routing mechanism balancing retrieval accuracy with computational efficiency.",
            "Two rooms and lights: Search a headline index or a full-article index; mix or expand when confidence is not HIGH.",
        ),
        (
            "Comparative evaluation: Benchmark the proposed dynamic routing approach against static threshold-based routing",
            "Comparative evaluation: Keep Phase 3B frozen (86% vs 84%). Add a frozen 40-query trap test (60% vs 20%) and dual-index P@5 on the same queries.",
        ),
        (
            "A dynamic SVM-based query routing model that is learnable at ceiling on development/cross-validation splits",
            "A dual-index SVM router: SHORT opens headlines, LONG opens full articles. Frozen Phase 3B (V2): 86.00% vs 84.00% word count. Frozen 40-query traps (twelve-feature model): 60.00% vs 20.00% word count. Development 100% CV is not the generalization claim.",
        ),
        (
            "An eight-feature semantic classifier engineered specifically for Urdu and Roman Urdu query characterization (canonical V2 order:",
            "A canonical eight-feature V2 extractor verified in Phase 3A, later extended with four intent flags (why / how / synthesis / fact cue) for trap-aware training. Phase 3A and Phase 3B remain the V2 frozen record.",
        ),
        (
            "A confidence-based three-tier routing architecture achieving an average confidence score of 98.18%",
            "Confidence lights: HIGH (≥85%) searches one room; MEDIUM (60–85%) hybrid-ranks both rooms; LOW (<60%) expands the query then hybrid-ranks.",
        ),
        (
            "An expanded Roman Urdu transliteration dictionary (30 → 179 words, a 6x growth)",
            f"An expanded Roman Urdu transliteration dictionary ({n_dict} on-disk pairs; earlier drafts cited 179), improving Roman Urdu retrieval coverage in the development P@15 experiment (92.5% on queries ULTRA scored 0.00%). That P@15 figure is supporting, not headline.",
        ),
        (
            "This thesis is scoped to the query-routing layer of an Urdu information retrieval pipeline.",
            "This thesis is scoped to query routing and to making that route change retrieval. It covers (i) SHORT versus LONG as headline-enough versus need-the-article, (ii) an SVM router with confidence lights, (iii) Roman Urdu transliteration, and (iv) two semantic indexes (headlines and full articles) over the same news corpus. It does not train a new encoder. Frozen tests are Phase 3B (50 labelled queries) and a later 40-query trap set. Dual-index P@5 uses depth-5 judgments, not P@15 as a judged metric.",
        ),
        (
            "Several adjacent problems are deliberately excluded from this scope",
            "Adjacent problems are excluded so the routing claim stays identifiable. The encoder, corpus preprocessing, and HNSW construction are inherited. Gains in Chapter 5 are therefore from routing and from opening a different room, not from a new embedding model. Query understanding beyond the engineered features (NER, syntax) is out of scope. The twelve-feature trap model is a later layer; it does not overwrite the frozen V2 Phase 3B table.",
        ),
        (
            "Chapter 3 — Mathematical Formulation. Presents the formal mathematical basis",
            "Chapter 3 — Mathematical Formulation. Presents the V2 feature vector, the SVM and lights, and cosine scoring in each room (headline matrix and full-article index).",
        ),
        (
            "Chapter 4 — Methods Developed. Describes the system architecture, 409-query training set",
            "Chapter 4 — Methods Developed. Describes two indexes, the SVM router, lights, 409-query V2 training, trap-augmented twelve-feature training, and frozen evaluation protocols.",
        ),
        (
            "Chapter 5 — Results and Discussion. Reports development-stage experiments and the later frozen Phase 3B",
            "Chapter 5 — Results and Discussion. Keeps development 100% CV separate from frozen 86/84, then reports 60/20 traps and dual-index P@5.",
        ),
        # --- Ch2 gap close ---
        (
            "Restating the gap analysis above in prose, three unresolved gaps carry directly into the design decisions",
            "Restating the gap in prose: first, ULTRA still uses one surface signal and one combined index, which motivates a learned room choice plus two rooms (Sections 3.1 and 3.6). Second, Roman Urdu work rarely sits inside retrieval, which motivates dictionary transliteration in Section 4.3. Third, no reviewed router combines a lightweight SVM with lights that mean one room / both rooms / expand-then-both, which motivates Section 3.5. Those gaps are what this thesis builds.",
        ),
        (
            "Most directly relevant to this thesis, Bashir, Qaiser, and Hussain (2026) proposed ULTRA",
            "Most directly relevant, Bashir, Qaiser, and Hussain (2026) proposed ULTRA: a dual-embedding Urdu architecture that already distinguished title/headline-level from full-content search using a length threshold, with precision gains above 90% on a large Urdu news corpus. ULTRA’s switch is still a static length tape. It does not learn headline-enough versus need-the-article, does not use script or why/how/fact cues, and does not turn confidence into mix-or-expand lights. This thesis keeps ULTRA’s two-level idea and replaces the tape with a learned dual-index router.",
        ),
        # --- Ch3 routing ---
        (
            "Each query q is converted into an eight-dimensional feature vector",
            "Each query q is converted into a feature vector. The frozen V2 path uses eight dimensions (Table 3.1). A later trap-augmented path appends four binary intent flags (has_causal, has_manner, has_synthesis, has_fact_cue) for a twelve-dimensional vector. Phase 3A and Phase 3B remain eight-feature V2. The twelve-feature model is trained later and tested on a separate frozen 40; it is not used to overwrite 86/84.",
        ),
        (
            "R(q) = Full Semantic Search,  if C(q) ≥ 85%  (HIGH)",
            "R(q) = one room (HEADLINE if SHORT, FULL_CONTENT if LONG),  if C(q) ≥ 85%  (HIGH)",
        ),
        (
            "R(q) = Hybrid Search,  if 60% ≤ C(q) < 85%  (MEDIUM)",
            "R(q) = hybrid mix of HEADLINE and FULL_CONTENT ranks,  if 60% ≤ C(q) < 85%  (MEDIUM)",
        ),
        (
            "R(q) = Query Expansion,  if C(q) < 60%  (LOW)",
            "R(q) = expand query, then hybrid mix of both rooms,  if C(q) < 60%  (LOW)",
        ),
        (
            "The specific thresholds 85% and 60% were chosen to reflect two qualitatively different levels of risk tolerance",
            "The 85% and 60% cuts are policy constants, not fitted on the frozen 40. HIGH means the SVM is trusted to open one room. MEDIUM means mix both rooms because the label may be wrong. LOW means the query is too thin, so a short expansion is appended and then both rooms are mixed. Lights change search, not the frozen 86/84 scoring rule. The frozen 40-query P@5 table compares one-room systems so rooms stay comparable.",
        ),
        (
            "All three routing tiers ultimately query the same underlying retrieval backend:",
            "HIGH confidence searches one room. MEDIUM and LOW mix two rooms built with the same encoder (paraphrase-multilingual-MiniLM-L12-v2): a headline embedding matrix and the full-article Chroma collection. Similarity in each room is cosine:",
        ),
        (
            "and the top-15 documents by sim(e_q, e_d) are returned. The HNSW (Hierarchical Navigable Small World) index approximates this nearest-neighbor search",
            "Each room returns a ranked list. HIGH uses one list. MEDIUM/LOW fuse the two lists with equal weight after min-max normalisation (LOW first appends a short expansion string). Development notebooks also report unjudged P@15. Frozen dual-index comparisons report graded P@5 because human (or protocol-confirmed) judgments stop at rank 5.",
        ),
        (
            "For retrieval quality, Precision at rank 15 (P@15) is used:",
            "Development retrieval quality is reported as P@15. Frozen dual-index quality is graded P@5 (Relevant = 1, Partially relevant = 0.5, Not relevant = 0):",
        ),
        (
            "P@15 = (Relevant documents in top-15 results) / 15",
            "P@5 = (sum of graded relevance in the top-5 results) / 5",
        ),
        (
            "To make Sections 3.1-3.5 concrete, this section traces a single Roman Urdu query",
            "To make Sections 3.1–3.5 concrete, this section traces “cricket match ka nateeja” (the result of the cricket match). Script detection flags Roman Urdu and dictionary-transliterates the query. V2 features are extracted (query_len = 4). The SVM typically returns SHORT with HIGH confidence, so only the headline room is searched. In the development retrieval experiment this query scored P@15 of 100.00% (Table 5.3). That example illustrates the HIGH-one-room path; it is not a frozen P@5 result.",
        ),
        # --- Ch4 ---
        (
            "This chapter describes the methodology developed to replace the static, length-based query routing mechanism of the base ULTRA framework",
            "This chapter describes how ULTRA's θ = 150 character tape is replaced as a whole: the decision (SVM SHORT/LONG) and the search (headline index versus full-article index), plus lights when the SVM is unsure. Roman Urdu is transliterated first. V2 (eight features, 409 queries) is the frozen Phase 3A/3B model. A later twelve-feature model adds 38 trap rows labelled by retrieval need; it is evaluated on a separate frozen 40.",
        ),
        (
            "Stage 1 — Language-variant normalization. An incoming query q is first passed through a Roman Urdu detection module.",
            f"Stage 1 — Language-variant normalization. Roman Urdu is detected and transliterated with the expanded dictionary ({n_dict} on-disk pairs) before Stage 2; native-script Urdu bypasses transliteration.",
        ),
        (
            "Stage 2 — Feature-based query representation.",
            "Stage 2 — Feature vector. V2 uses the eight Phase 3A features. The trap-augmented model appends four binary intent flags (causal, manner, synthesis, fact cue). Neither vector uses θ = 150 as an input feature.",
        ),
        (
            "Stage 4 — Tiered retrieval execution.",
            "Stage 4 — Two rooms and lights. HIGH: search only HEADLINE or only FULL_CONTENT. MEDIUM: hybrid both rooms. LOW: expand, then hybrid. This is the operational replacement of ULTRA's single combined-text search.",
        ),
        (
            "HIGH confidence (≥ 85%): the query is routed directly to full semantic search against the ChromaDB index.",
            "HIGH (≥ 85%): one semantic room, the room the SVM named. MEDIUM (60–85%): both semantic rooms, equal hybrid. LOW (< 60%): append a short expansion (Urdu تفصیل وجوہات or English “detail reasons news”), then hybrid. Lights do not change the frozen 86/84 table; they only change which lists are merged at search time.",
        ),
        (
            "The retrieval backend is built over a corpus of 111,860 Urdu news articles",
            "The retrieval backend uses one news corpus (about 111,860 Urdu articles) and two rooms: a headline embedding cache and a full-article Chroma collection (urdu_news), both encoded with the same multilingual MiniLM. HIGH searches one room. MEDIUM and LOW mix both. This is not “the same combined Chroma for every tier.”",
        ),
        (
            "Detected Roman Urdu queries are passed through a transliteration dictionary that was expanded from an initial 30 word-pairs to 179 word-pairs",
            f"Detected Roman Urdu queries are passed through a transliteration dictionary expanded from about 30 pairs to {n_dict} on-disk pairs (earlier drafts cited 179). Coverage is still finite and token-level; proper nouns and loanwords remain the main residual error (Section 5.8).",
        ),
        (
            "Because Roman Urdu lacks standardized spelling conventions, a dedicated detection step identifies queries written in Latin-script transliterated Urdu prior to feature extraction. Detected Roman Urdu queries are passed through a transliteration dictionary that was expanded from an initial 30 word-pairs to 179 word-pairs",
            f"Because Roman Urdu lacks standardized spelling, Latin-script queries are detected before features are extracted and then passed through the expanded dictionary ({n_dict} on-disk pairs).",
        ),
        (
            "The classifier was developed iteratively: an initial illustrative training run validated an eight-feature RBF-kernel SVM design",
            "The classifier was developed in layers. V2 is an eight-feature RBF SVM trained on 409 queries after gap-fill and label audit; Phase 3A/3B freeze that object. A later trap-augmented twelve-feature SVM was trained on 447 queries (409 plus 38 traps; T033 and T036 were already in the 409). The deployed pickle is that twelve-feature model. Phase 3B 86/84 was not overwritten. H001–H040 were never used for training.",
        ),
        (
            "Deliberately, no exhaustive grid search over C and γ was performed",
            "No exhaustive grid search over C and γ was performed; defaults were retained and sensitivity checked (Section 5.7). Development accuracy near 100% is not the frozen result. Frozen Phase 3B V2 is 86.00%, not 100%. The later twelve-feature model’s 96% on overlapping Phase 3B rows is leakage and is not reported as a result.",
        ),
        (
            "HIGH confidence (≥ 85%): the query is routed directly to full semantic search against the ChromaDB index. MEDIUM confidence",
            "HIGH (≥ 85%): one room. MEDIUM (60–85%): hybrid both rooms. LOW (< 60%): expand, then hybrid. Effort scales with uncertainty, but the rooms are two semantic indexes, not BM25-plus-dense on one combined text store.",
        ),
        (
            "After the V2 SVM and scaler were trained on the 409-query set and persisted",
            "After V2 was trained on 409 queries and persisted, Phase 3A verified the eight-feature path and Phase 3B scored the frozen 60-query file without refitting. A later twelve-feature pickle replaced the deployed files for trap-aware routing; V2 remains backed up and is the Phase 3B 86/84 model. A different 40-query sheet (H001–H040) was then frozen and never trained on.",
        ),
        (
            "Phase 2.5 does not train or modify the classifier. It tests whether the SHORT/LONG distinction near 5-6 words matches human retrieval need.",
            "Phase 2.5 does not train the classifier. It tests whether SHORT/LONG near five–six words matches retrieval need, using headline versus full-article semantic search and rank-1–5 human labels (330 rows). A later frozen 40-query dual-index study (Section 5.17) adds 400 more depth-5 judgments on traps. Because judgments stop at rank 5, P@15 from these pilots is not a judged metric.",
        ),
        # --- Ch5 ---
        (
            "This chapter reports the empirical evaluation of the proposed dynamic routing framework",
            "This chapter keeps three evaluation layers apart. Sections 5.1–5.12 are development: they show the task is learnable (including 100% on some splits) and that Roman Urdu search is no longer zero. Those figures are not the generalization claim. Section 5.13 is frozen V2 Phase 3B: 86.00% versus 84.00% word count. Sections 5.16–5.17 are a later frozen 40-query trap study: 60% versus 20% on need labels, and dual-index P@5 that does not yet favour the SVM. Do not replace 86/84 with 100% or with 96%.",
        ),
        (
            "Table 5.1 compares the static character-length threshold baseline (θ = 150) against the proposed dynamic classifiers",
            "Table 5.1 compares θ = 150 against the development classifiers on the train/test split used while building the model. Those figures show learnability. They are not frozen generalization. Frozen V2 is Table 5.8 / Section 5.13 (86.00% vs 84.00%). Frozen traps are Section 5.16 (60.00% vs 20.00%).",
        ),
        (
            "On this development split, the dynamic classifiers achieve a 50 percentage-point improvement",
            "On this development split, dynamic classifiers gain 50 points over θ = 150 (100% vs 50%). That gap shows θ = 150 is the wrong cut for this labelled set (Figure 4.2). It is not the held-out margin. Frozen V2 is two points above a six-word rule and not significant. Frozen traps beat word count on labels, then lose on P@5.",
        ),
        (
            "Strengths. First, routing accuracy is both very high and independently confirmed across three separate evaluation conditions",
            "Strengths. First, the routing task is learnable, and frozen V2 still generalizes to 86% rather than collapsing to chance. Second, sub-millisecond local inference is a real deployment fact against LLM-style routers. Third, Roman Urdu is no longer a zero-P@15 dead end in the development retrieval experiment. Fourth, trap labels show that a cue-aware SVM can beat a six-word tape on need (60% vs 20%), which θ = 150 cannot.",
        ),
        (
            "Weaknesses. First, as detailed in Section 5.8, the Roman Urdu transliteration module's dictionary-based design has a demonstrated failure mode",
            "Weaknesses. First, dictionary transliteration still fails on loanwords and proper nouns. Second, frozen V2 only just matches word count (86/84; p = 1.0). Third, trap gains are cue-word gains, not a general “need” detector (27.27% without cues). Fourth, dual-index P@5 on those traps does not beat word count (33.00% vs 36.50%). Fifth, development 100% and AUC = 1.000 must not be cited as the paper result.",
        ),
        (
            "Taken together, these results support four main conclusions. First, replacing ULTRA's static character-length threshold",
            "Taken together, the honest story is this. θ = 150 never fires LONG on the Phase 2.5 and trap queries used here, so it is the wrong switch. A six-word tape is a much stronger simple rule. Frozen V2 only just matches it (86% vs 84%). Trap labels that mean headline-enough versus need-the-article let a cue-aware SVM beat that tape on classification (60% vs 20%), but only where why/how/fact words appear. Dual-index P@5 on those same traps does not yet beat word count. Lights are implemented; the frozen P@5 table used one-room systems so rooms stayed comparable.",
        ),
        (
            "Fourth, the feature importance and ablation results together reveal a nuanced picture",
            "Fourth, development permutation/ablation numbers describe an earlier feature set, not frozen V2, and must not be used to claim that question-words drive the deployed model. V2 is script and length; V3 adds cue flags. Collectively these findings answer RQ1–RQ3 with the limits stated in Chapter 6.",
        ),
        # --- Ch6 ---
        (
            "This thesis addressed a specific limitation of the ULTRA framework for Urdu information retrieval: its reliance on a single, static, character-length threshold",
            "This thesis addressed ULTRA as a whole: not only the θ = 150 decision, but the search that must follow it. SHORT now opens a headline semantic room; LONG opens a full-article room. Lights mix the rooms when the SVM is unsure. Roman Urdu is dictionary-transliterated first.",
        ),
        (
            "On development/cross-validation data, the proposed dynamic routing mechanism raised routing accuracy from 50.00%",
            "Headline numbers: frozen Phase 3B 86.00% (V2 SVM) vs 84.00% (six-word rule), McNemar p = 1.0000. Frozen 40-query traps 60.00% vs 20.00% vs 50.00% (θ = 150), McNemar 16–0. Cue split: 100% with V3 flags, 27.27% without (tied with word count). Dual-index graded P@5 on those 40: word count 36.50%, always headlines / θ = 150 35.00%, always full 34.25%, SVM 33.00%. nDCG@5 on the same sheet is highest for always-headline / θ = 150 (0.6868), then word count (0.6476), SVM (0.6149), always-full (0.6020). Development 100% CV and 92.5% Roman P@15 remain supporting, not headline, results.",
        ),
        (
            "These results answer RQ1-RQ3 with different degrees of confidence.",
            "RQ1 is supported with a limit: script and length features plus why/how/fact flags predict room choice when those flags fire. RQ2 is yes on trap classification, only a 2-point frozen V2 edge versus word count, and not yet on trap P@5. RQ3 is implemented (HIGH/MEDIUM/LOW lights) and shown on live examples; it was not the scoring rule for the 40-query P@5 table, which used one-room SVM versus one-room baselines.",
        ),
        (
            "A dynamic SVM-based query routing model (development: 100% vs. 50% θ = 150; frozen Phase 3B: 86.00% vs. 84.00% word-count).",
            "Two-room SVM routing, frozen 86/84 (V2) and frozen 60/20 (traps). Substantiated by Sections 5.13 and 5.16.",
        ),
        (
            "An eight-feature semantic classifier engineered specifically for Urdu and Roman Urdu query characterization (canonical V2 order: urdu_ratio, roman_ratio, has_urdu, has_roman, query_len, char_len, mixed, urdu_chars), verified against the deployed scaler and SVM in Phase 3A.",
            "Phase 3A eight-feature V2 path, plus a later twelve-feature trap model. Phase 3B was not retrained away.",
        ),
        (
            "A confidence-based three-tier routing architecture (98.18% average confidence).",
            "HIGH / MEDIUM / LOW lights that choose one room, both rooms, or expand-then-both (Section 3.5, 4.6).",
        ),
        (
            "An expanded Roman Urdu transliteration dictionary (30 → 179 words), raising Roman Urdu P@15 to 92.5%.",
            f"An expanded Roman Urdu dictionary ({n_dict} on-disk pairs), with development Roman P@15 of 92.5% as a supporting result (Section 5.2).",
        ),
        (
            "An independent frozen Phase 3B evaluation (60 queries; 50 primary) of the deployed V2 SVM against a frozen word-count baseline, with Phase 3A verification of the eight-feature inference path.",
            "Frozen Phase 3B (50 primary) plus a frozen 40-query trap set and 400 dual-index relevance rows (Sections 5.16–5.17).",
        ),
        (
            "Untested confidence tiers. As noted in Section 6.1, the MEDIUM and LOW confidence tiers are implemented and formally specified",
            "Lights versus P@5 protocol. MEDIUM and LOW are implemented in code and shown on live examples. The frozen 40-query P@5 table compares one-room systems so that headline versus full lists remain comparable. Hybrid P@5 is therefore not claimed as a frozen number.",
        ),
        (
            "Untested confidence tiers. As noted in Section 6.1, the MEDIUM and LOW confidence tiers",
            "Lights versus P@5 protocol. MEDIUM and LOW are implemented in code. The frozen 40-query P@5 table compares one-room systems. Hybrid P@5 is not a frozen number.",
        ),
        (
            "This thesis set out to answer a narrow question — can a learned, multi-feature router outperform a static threshold for Urdu query routing",
            "The question was whether Urdu search should still flip rooms with 150 letters. It should not. A six-word tape is closer to the truth but fails when a short query needs a story or a long query needs one fact. This thesis built the two rooms, taught an SVM to choose, and added lights. Frozen V2 almost ties word count (86/84). Frozen traps beat word count on labels (60/20) only with cue words. Dual-index P@5 on those traps does not yet beat the tape. That is a complete, defensible, and IEEE-honest story: the switch is right; the remaining work is need features that are not a keyword list, and a corpus that actually answers why.",
        ),
        (
            "Dedicated stress-testing of the MEDIUM and LOW confidence tiers.",
            "Need features beyond cue words, and a corpus that answers why. The cue split shows the twelve-feature SVM is still a keyword detector on traps. Dual-index P@5 shows the right room still retrieves “price rose” instead of causes. Future work should learn need without a closed cue list, and should measure hybrid-light P@5 only after that retrieval gap narrows.",
        ),
    ]

    hit, miss = replace_prefixed(doc, jobs)

    # Insert 5.16 / 5.17 after Phase 2.5 closing paragraph
    anchor = None
    for p in doc.paragraphs:
        t = p.text or ""
        if t.startswith("A separate 8-row LLM-judging smoke test") or t.startswith(
            "A separate 8-row LLM-judging smoke test"
        ):
            anchor = p
            break
        if t.startswith("Human relevance labels on the 330 judged rows") or t.startswith(
            "Human relevance labels on the 330 judged rows"
        ):
            anchor = p
    if anchor is None:
        for p in doc.paragraphs:
            if (p.text or "").startswith("5.15 Human Retrieval Validation Results"):
                anchor = p

    h2_name = "Heading 2"
    for p in doc.paragraphs:
        if (p.text or "").startswith("5.15"):
            try:
                h2_name = p.style.name
            except Exception:
                pass
            break

    if anchor is not None:
        blocks = [
            (h2_name, "5.16 Frozen held-out trap classification (n = 40)"),
            (
                "Normal",
                "After the V2 Phase 3B file was frozen, 38 additional trap queries were labelled with the headline-enough rule and used only for a later twelve-feature SVM. A different 40-query sheet (H001–H040) was then frozen and never trained on. Labels on that sheet were produced with a written protocol (headline enough versus need the article), assisted in a first pass, then saved by the author; they are not two independent human raters. Agreement with the protocol is 40/40. Accuracy: SVM 60.00%, six-word rule 20.00%, θ = 150 50.00%. McNemar versus word count: 16 SVM-only-correct, 0 word-count-only-correct, exact p < 0.001. By construction, θ = 150 predicts SHORT for every query in this set (all are under 150 characters).",
            ),
            (
                "Normal",
                "The cue split is the scientific limit, not a footnote. On the 18 queries where a V3 intent flag fires (causal, manner, synthesis, or fact phrasing), SVM accuracy is 18/18 and word-count accuracy is 2/18 (11.11%). On the 22 queries without those flags (for example ٹیم ہار کا تجزیہ or آج کراچی میں زیادہ سے زیادہ حرارت), both systems score 27.27%. The router therefore beats length when the query wears why/how/fact words; otherwise it still behaves like the six-word tape. IEEE writing must state that split. It must not replace Phase 3B 86% with 60% or with 100%.",
            ),
            (h2_name, "5.17 Frozen held-out dual-index P@5 (400 judgments)"),
            (
                "Normal",
                "The same 40 queries were searched in both rooms (top-5 each). Four hundred headline rows were labelled Relevant / Partially relevant / Not relevant (61 / 155 / 184). Graded P@5 (1.0 / 0.5 / 0.0) was scored for one-room systems so that Phase-style lists stay comparable: word-count router 36.50%, always-headline and θ = 150 35.00%, always-full 34.25%, SVM 33.00%. nDCG@5 does not follow that P@5 order: always-headline / θ = 150 0.6868, word count 0.6476, SVM 0.6149, always-full 0.6020. So the classification win on traps does not appear as a retrieval win. Typical failure: a LONG why-query still retrieves “petrol became expensive” rather than causes, in both rooms. Right room, weak articles — that is the remaining ULTRA claim, not a licence to hide the table.",
            ),
            (
                "Normal",
                "This P@5 comparison used the deployed twelve-feature SVM as a one-room router (HIGH path only). It is not a frozen hybrid-light evaluation. Phase 2.5 (33 queries, depth 5) remains a separate earlier pilot and must not be averaged with these 40.",
            ),
        ]
        for style, text in reversed(blocks):
            insert_after(anchor, text, style)

    for p in doc.paragraphs:
        if (p.text or "").strip() in (
            "5.15 Human Retrieval Validation Results (Phase 2.5)",
            "5.15 Human Retrieval Validation Results (Phase 2.5)",
        ):
            insert_after(p, "5.17 Frozen held-out dual-index P@5", p.style.name)
            insert_after(p, "5.16 Frozen held-out trap classification", p.style.name)
            break

    # Light table-cell fixes
    cell_map = {
        "Full Semantic Search": "One semantic room (headline or full article)",
        "full semantic search": "one semantic room",
        "Hybrid Search": "Hybrid mix of both rooms",
        "Query Expansion": "Expand, then hybrid both rooms",
    }
    n_cells = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text
                for old, new in cell_map.items():
                    if old in t:
                        cell.text = t.replace(old, new)
                        n_cells += 1

    doc.save(str(SRC))
    print("dict_count", n_dict)
    print("replaced", len(hit), "missed", len(miss))
    for m in miss:
        print("MISS:", m)
    print("table_cells", n_cells)
    print("saved", SRC)
    print("backup", BAK)


if __name__ == "__main__":
    main()
