# -*- coding: utf-8 -*-
"""Read-only dump of live thesis for consistency audit."""
from pathlib import Path
from docx import Document

HERE = Path(__file__).resolve().parent
SRC = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"
OUT = HERE / "_audit_dump.txt"

d = Document(str(SRC))
lines = []
lines.append(f"paragraphs={len(d.paragraphs)} tables={len(d.tables)}\n")
for i, p in enumerate(d.paragraphs):
    t = p.text or ""
    if not t.strip():
        continue
    style = p.style.name if p.style else ""
    lines.append(f"[{i}|{style}] {t}\n")
OUT.write_text("".join(lines), encoding="utf-8")
print("wrote", str(OUT), "n", len(lines))
