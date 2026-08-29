# -*- coding: utf-8 -*-
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
SRC = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"
LOG = HERE / "_audit_changes.txt"

HEADS = (
    "5.18 Official frozen retrieval system (M0)",
    "5.19 Phase 2 development/validation known-item (n = 78)",
    "5.20 Phase 11 ablation: M0 remains official",
    "5.21 Phase 12 new known-item evaluation (K001–K040)",
    "5.22 Phase 12 naturalistic human evaluation (U001–U040)",
    "5.23 Final comparison of evaluation settings",
    "5.24 Discussion of the official IR results",
)


def main() -> None:
    doc = Document(str(SRC))
    src_h2 = None
    for p in doc.paragraphs:
        if (p.text or "").strip() == HEADS[0]:
            src_h2 = p
            break
    if src_h2 is None:
        raise SystemExit("5.18 heading missing")
    src_pr = src_h2._p.find(qn("w:pPr"))
    n = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t in HEADS:
            dest = p._p
            old = dest.find(qn("w:pPr"))
            if old is not None:
                dest.remove(old)
            if src_pr is not None:
                dest.insert(0, deepcopy(src_pr))
            try:
                p.style = "Heading 2"
            except Exception:
                pass
            n += 1
    notes = ["Forced Heading 2 on %d official M0 section headings so TOC can pick up 5.18-5.24" % n]
    styles = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t in HEADS or t.startswith("Table ") and "official M0" in t or t.startswith("Table ") and "Do not average rows" in t:
            styles.append("%s | %s" % ((p.style.name if p.style else ""), t[:90]))
        if (p.text or "").startswith("H001–H040 human Success@5"):
            styles.append("%s | %s" % ((p.style.name if p.style else ""), (p.text or "")[:90]))
        if "Script-wise U Success@5" in (p.text or ""):
            styles.append("%s | %s" % ((p.style.name if p.style else ""), (p.text or "")[:90]))
    try:
        doc.save(str(SRC))
        notes.append("Saved live docx")
    except PermissionError:
        alt = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA_audit.docx"
        doc.save(str(alt))
        notes.append("locked; saved " + alt.name)
    Path(HERE / "_audit_verify3.txt").write_text("\n".join(notes + styles), encoding="utf-8")
    prev = LOG.read_text(encoding="utf-8") if LOG.exists() else ""
    LOG.write_text(prev + "\n".join("- " + x for x in notes) + "\n", encoding="utf-8")
    print("ok", n)


if __name__ == "__main__":
    main()
