# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document

HERE = Path(__file__).resolve().parent
SRC = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"
LOG = HERE / "_audit_changes.txt"
NEEDLE = "Contributions 1 through 3 are architectural: together, they constitute the dynamic routing system itself"
PREFIX = (
    "Contribution 1 is the official frozen retriever M0 (script-aware BM25). "
    "SVM SHORT/LONG rooms and confidence lights are historical Layer A, not the official IR system."
)


def main() -> None:
    doc = Document(str(SRC))
    note = "UNCHANGED: Ch.1 contributions prose not found"
    for p in doc.paragraphs:
        t = p.text or ""
        if NEEDLE in t:
            if PREFIX[:40] in t:
                note = "SKIP already labeled: Ch.1 contributions prose"
                break
            if p.runs:
                p.runs[0].text = PREFIX + " " + t
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(PREFIX + " " + t)
            note = "Ch.1 contributions prose: official system is M0; SVM rooms/lights are Layer A"
            break
    doc.save(str(SRC))
    prev = LOG.read_text(encoding="utf-8") if LOG.exists() else ""
    LOG.write_text(prev + "- " + note + "\n", encoding="utf-8")
    print(note)


if __name__ == "__main__":
    main()
