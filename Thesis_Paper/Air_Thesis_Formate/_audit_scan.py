# -*- coding: utf-8 -*-
"""Scan remaining presentation contradictions. Read-only."""
from pathlib import Path
from docx import Document

HERE = Path(__file__).resolve().parent
d = Document(str(HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"))
needles = [
    "the frozen generalization test is Phase 3B",
    "replaces the tape with a learned dual-index router",
    "Those gaps are what this thesis builds",
    "cosine-similarity scoring function used by the retrieval backend",
    "proposed dynamic query routing system",
    "the same combined Chroma for every tier",
    "80%",
    "real-world accuracy",
    "average",
    "M1 replaced",
    "official M0" ,
]
# more targeted scans
flags = []
for i, p in enumerate(d.paragraphs):
    t = p.text or ""
    st = (p.style.name if p.style else "")
    if st.lower().startswith("toc") or "table of" in st.lower():
        continue
    checks = [
        ("Phase 3B as THE frozen IR test", "frozen generalization test is Phase 3B" in t),
        ("unqualified dual-index as thesis IR", "replaces the tape with a learned dual-index router" in t and "Official M0" not in t),
        ("80 percent usefulness", "80%" in t or "80.00%" in t),
        ("M1 replaced M0", "M1 replaced" in t or "replaced M0" in t.lower() and "not replaced" not in t.lower()),
        ("average the three", "average" in t.lower() and "87.18" in t),
        ("87.18 as unseen", "87.18" in t and ("unseen usefulness" in t.lower() or "real-world accuracy" in t) and "not" not in t.lower()[: t.lower().find("87.18")+80] if "87.18" in t else False),
    ]
    for label, hit in checks:
        if hit:
            flags.append("[%d|%s] %s :: %s" % (i, st[:20], label, t[:160].replace("\n", " ")))

# print heading sequence 4.11-6.1
heads = []
for i, p in enumerate(d.paragraphs):
    st = p.style.name if p.style else ""
    t = (p.text or "").strip()
    if st.startswith("Heading") and t[:4] in ("4.11", "4.12", "4.13", "5.1 ", "5.17", "5.18", "5.19", "5.20", "5.21", "5.22", "5.23", "5.24", "Chap", "6.1 ", "6.2 ", "6.6"):
        heads.append("[%d|%s] %s" % (i, st, t[:120]))
    elif st.startswith("Heading") and (t.startswith("5.") or t.startswith("4.1") or t.startswith("Chapter")):
        if any(t.startswith(x) for x in ("4.11", "4.12", "4.13", "5.2 ", "5.10", "5.13", "5.17", "5.18", "Chapter 5", "Chapter 6", "6.1", "6.2", "6.6", "3.6")):
            heads.append("[%d|%s] %s" % (i, st, t[:120]))

out = ["HEADINGS"] + heads + ["", "FLAGS %d" % len(flags)] + flags
# caption samples
caps = []
for p in d.paragraphs:
    if (p.style.name if p.style else "") == "Caption":
        caps.append((p.text or "")[:140])
out.append("")
out.append("CAPTIONS %d" % len(caps))
out.extend(caps)
Path(HERE / "_audit_scan.txt").write_text("\n".join(out), encoding="utf-8")
print("flags", len(flags), "captions", len(caps), "heads", len(heads))
