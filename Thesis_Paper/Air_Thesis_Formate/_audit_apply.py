# -*- coding: utf-8 -*-
"""Consistency audit of the live AU Word thesis. Presentation only. No experiments."""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

HERE = Path(__file__).resolve().parent
SRC = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"
BAK = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.pre_consistency_audit.bak.docx"
LOG = HERE / "_audit_changes.txt"

CAPTION_RE = re.compile(r"^(Figure|Table)\s+[\d.]+\.\s*(.*)$", re.S)
TOC_PREFIX = re.compile(
    r"^5\.17 Frozen held-out dual-index P@5 \(400 judgments\)\s*\t?\d*",
)

HEADING_MAP = {
    "5.18 Official frozen retrieval system (M0)": "5.18 Official frozen retrieval system (M0)",
    "5.19 Phase 2 development/validation known-item (n = 78)": "5.19 Phase 2 development/validation known-item (n = 78)",
    "5.20 Phase 11 ablation: M0 remains official": "5.20 Phase 11 ablation: M0 remains official",
    "5.21 Phase 12 new known-item evaluation (K001–K040)": "5.21 Phase 12 new known-item evaluation (K001–K040)",
    "5.22 Phase 12 naturalistic human evaluation (U001–U040)": "5.22 Phase 12 naturalistic human evaluation (U001–U040)",
    "5.23 Final comparison of evaluation settings": "5.23 Final comparison of evaluation settings",
    "5.24 Discussion of the official IR results": "5.24 Discussion of the official IR results",
}

TABLE_RENAMES = (
    ("Table 3. Phase 2 n=78", "Table 22. Phase 2 n=78"),
    ("Table 4. Phase 11 M0–M4", "Table 23. Phase 11 M0–M4"),
    ("Table 5. Phase 12 K ExactSource", "Table 24. Phase 12 K ExactSource"),
    ("Table 6. Phase 12 U human", "Table 25. Phase 12 U human"),
    ("Table 7. Script-wise U Success@5", "Table 26. Script-wise U Success@5"),
    ("Table 8. Do not average rows.", "Table 27. Do not average rows."),
)


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def try_style(p, name: str) -> None:
    try:
        p.style = name
    except Exception:
        pass


def append_field(paragraph: Paragraph, instruction: str, placeholder: str) -> None:
    def add_child(child) -> None:
        run = OxmlElement("w:r")
        run.append(child)
        paragraph._p.append(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    add_child(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction.strip()} "
    add_child(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    add_child(sep)
    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = placeholder
    add_child(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    add_child(end)


def clear_runs(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag == qn("w:r"):
            paragraph._p.remove(child)


def delete_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def insert_empty_after(paragraph: Paragraph) -> Paragraph:
    new_el = deepcopy(paragraph._p)
    paragraph._p.addnext(new_el)
    np = Paragraph(new_el, paragraph._parent)
    try_style(np, "Normal")
    clear_runs(np)
    return np


def insert_after(paragraph, text: str, style_name: str):
    new_el = deepcopy(paragraph._p)
    paragraph._p.addnext(new_el)
    np = Paragraph(new_el, paragraph._parent)
    try_style(np, style_name)
    set_text(np, text)
    return np


def enable_update_on_open(doc: Document) -> None:
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)
    else:
        existing.set(qn("w:val"), "true")


def delete_until(doc: Document, start_title: str, end_title: str) -> Paragraph:
    paras = list(doc.paragraphs)
    start_i = end_i = None
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if t == start_title and start_i is None:
            start_i = i
        elif start_i is not None and t == end_title:
            end_i = i
            break
    if start_i is None or end_i is None:
        raise SystemExit(f"could not find block {start_title!r} .. {end_title!r}")
    for p in paras[start_i + 1 : end_i]:
        delete_paragraph(p)
    return paras[start_i]


def clean_m0_text(t: str) -> str:
    t = TOC_PREFIX.sub("", t).strip()
    for old, new in TABLE_RENAMES:
        t = t.replace(old, new)
    for key, clean in HEADING_MAP.items():
        if key in t and not t.startswith("Table"):
            return clean
    return t


def style_m0_para(p) -> None:
    t = (p.text or "").strip()
    if t.startswith("5.1") and t[:4] in ("5.18", "5.19", "5.20", "5.21", "5.22", "5.23", "5.24"):
        try_style(p, "Heading 2")
    elif t.startswith("Table 2") and t[6:8].isdigit():
        try_style(p, "Caption")


def move_m0_blocks(doc: Document, changes: list[str]) -> int:
    paras = list(doc.paragraphs)
    start = end = None
    for i, p in enumerate(paras):
        t = p.text or ""
        if "5.18 Official frozen retrieval system (M0)" in t and start is None:
            start = i
        if start is not None and "Do not retune M0 on K, U, or H001–H040" in t:
            end = i
            break
    if start is None or end is None:
        changes.append("MISS: could not locate §§5.18–5.24 block")
        return 0

    body_ch6 = None
    for p in paras:
        if (p.text or "").strip() == "Chapter 6: Conclusions and Recommendations":
            style = p.style.name if p.style else ""
            if style.startswith("Heading"):
                body_ch6 = p
    if body_ch6 is None:
        raise SystemExit("body Chapter 6 heading not found")

    already_before = (end + 1 < len(paras)) and (paras[end + 1] is body_ch6)
    moved_paras = [paras[i] for i in range(start, end + 1)]
    if not already_before:
        els = [p._p for p in moved_paras]
        target = body_ch6._p
        for el in els:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
            target.addprevious(el)
        changes.append(
            f"MOVED §§5.18–5.24 from TOC region (old paras {start}-{end}) to immediately before Chapter 6 body"
        )
    else:
        changes.append(f"§§5.18–5.24 already immediately before Chapter 6 (paras {start}-{end})")

    for p in moved_paras:
        set_text(p, clean_m0_text(p.text or ""))
        style_m0_para(p)
    changes.append("Cleaned TOC-field prefixes from §§5.18–5.24 headings and body")
    changes.append("Renamed official result tables Table 3–8 → Table 22–27 to avoid clashing with existing tables")
    return len(moved_paras)


def is_front_matter(p) -> bool:
    s = (p.style.name if p.style else "").lower()
    return s.startswith("toc") or "table of" in s


def replace_if(doc, needle: str, new: str, note: str, changes: list[str], whole: bool = True) -> None:
    for p in doc.paragraphs:
        if is_front_matter(p):
            continue
        t = p.text or ""
        if needle in t:
            if whole:
                if t == new or t.startswith(new[:80]):
                    changes.append("SKIP already applied: " + note)
                    return
                set_text(p, new)
            else:
                if new.split(" ", 3)[0] in t[:120] and "Layer A" in t[:200]:
                    changes.append("SKIP already labeled: " + note)
                    return
                set_text(p, t.replace(needle, new, 1))
            changes.append(note)
            return
    changes.append("UNCHANGED (not found): " + note)


def prepend_if(doc, needle: str, prefix: str, note: str, changes: list[str]) -> None:
    for p in doc.paragraphs:
        if is_front_matter(p):
            continue
        t = p.text or ""
        if needle in t:
            if prefix[:50] in t:
                changes.append("SKIP already labeled: " + note)
                return
            set_text(p, prefix.rstrip() + " " + t)
            changes.append(note)
            return
    changes.append("UNCHANGED (not found): " + note)


def add_seq_captions(doc: Document) -> tuple[int, int]:
    n_fig = n_tab = 0
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if style != "Caption":
            continue
        raw = (p.text or "").strip()
        m = CAPTION_RE.match(raw)
        if not m:
            continue
        kind, title = m.group(1), m.group(2).strip()
        clear_runs(p)
        p.add_run(f"{kind} ")
        append_field(p, f"SEQ {kind} \\* ARABIC", "1")
        p.add_run(f". {title}")
        if kind == "Figure":
            n_fig += 1
        else:
            n_tab += 1
    return n_fig, n_tab


def rebuild_toc(doc: Document, changes: list[str]) -> None:
    toc_h = delete_until(doc, "Table of Contents", "List of Figures")
    toc_p = insert_empty_after(toc_h)
    append_field(
        toc_p,
        'TOC \\o "1-3" \\h \\z \\u',
        "Open this file in desktop Word and click Yes when asked to update fields.",
    )
    fig_h = delete_until(doc, "List of Figures", "List of Tables")
    fig_p = insert_empty_after(fig_h)
    append_field(
        fig_p,
        'TOC \\c "Figure" \\h \\z',
        "Open this file in desktop Word and click Yes when asked to update fields.",
    )
    tab_h = delete_until(doc, "List of Tables", "Symbols and Abbreviations")
    tab_p = insert_empty_after(tab_h)
    append_field(
        tab_p,
        'TOC \\c "Table" \\h \\z',
        "Open this file in desktop Word and click Yes when asked to update fields.",
    )
    enable_update_on_open(doc)
    changes.append(
        "Replaced typed Table of Contents, List of Figures, and List of Tables with Word TOC/SEQ fields; set updateFields on open"
    )


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"missing {SRC}")
    shutil.copy2(SRC, BAK)
    changes: list[str] = [f"Safety copy: {BAK.name}"]
    doc = Document(str(SRC))

    n = move_m0_blocks(doc, changes)
    changes.append(f"M0 block paragraphs processed: {n}")

    # Headings / captions: exact replace
    replace_if(
        doc,
        "Figure 1. System architecture: Roman Urdu dictionary, SVM room choice (headline vs full article), and HIGH/MEDIUM/LOW lights.",
        "Figure 1. Historical Layer A architecture: Roman Urdu dictionary, SVM room choice (headline vs full article), and HIGH/MEDIUM/LOW lights. Not the official M0 BM25 pipeline.",
        "Figure 1 caption: labeled historical Layer A, not M0",
        changes,
    )
    replace_if(
        doc,
        "Figure 4. Performance comparison: proposed system vs. static and LLM-style baselines.",
        "Figure 4. Layer A performance comparison: SVM router vs. static and LLM-style baselines (not official M0).",
        "Figure 4 caption: labeled Layer A, not M0",
        changes,
    )
    replace_if(
        doc,
        "Figure 15. Diagnostic external validation on 50 unseen native-Urdu queries (development/robustness layer). The deployed V2 model records 98.00% on this 50-query diagnostic set (training_info.json); the frozen generalization test is Phase 3B (Section 5.13).",
        "Figure 15. Layer A diagnostic: 50 native-Urdu queries, V2 routing accuracy 98.00%. This is SVM classification, not official M0 ExactSource Hit@5. Official IR tests are Sections 5.19–5.22.",
        "Figure 15 caption: SVM diagnostic, not official M0 / not unseen IR usefulness",
        changes,
    )
    replace_if(
        doc,
        "Figure 18. Three evaluation layers that must not be mixed: development/CV, frozen Phase 3B (86/84), and frozen traps (60/20/50).",
        "Figure 18. Layer A evaluation layers that must not be mixed with official M0: development/CV, frozen Phase 3B SVM (86/84), and frozen traps (60/20/50). Official IR is Sections 5.18–5.24.",
        "Figure 18 caption: Layer A only, not official M0",
        changes,
    )
    replace_if(
        doc,
        "Table 16. Retrieval precision at rank 15 (P@15).",
        "Table 16. Development-only retrieval precision at rank 15 (P@15). Not official M0.",
        "Table 16 caption: development-only / not M0",
        changes,
    )
    replace_if(
        doc,
        "4.11 Independent Frozen Evaluation Protocol (Phases 3A–3B)",
        "4.11 Independent Frozen Evaluation Protocol (Phases 3A–3B, Layer A SVM)",
        "§4.11 heading: labeled Layer A SVM",
        changes,
    )
    replace_if(
        doc,
        "5.2 Retrieval Precision (P@15)",
        "5.2 Retrieval Precision (P@15, Layer A development experiment)",
        "§5.2 heading: Layer A development experiment",
        changes,
    )
    replace_if(
        doc,
        "5.10 Strengths and Weaknesses of the Proposed System",
        "5.10 Strengths and Weaknesses of Layer A (SVM / MiniLM)",
        "§5.10 heading: Layer A, not official M0",
        changes,
    )
    replace_if(
        doc,
        "5.13 Frozen Phase 3B Generalization Results",
        "5.13 Frozen Phase 3B SVM classification results (Layer A)",
        "§5.13 heading: SVM classification Layer A, not official IR",
        changes,
    )
    replace_if(
        doc,
        "5.17 Frozen held-out dual-index P@5 (400 judgments)",
        "5.17 Frozen held-out dual-index P@5 (400 judgments, Layer A MiniLM; not official M0)",
        "§5.17 heading: Layer A MiniLM, not official M0",
        changes,
    )
    replace_if(
        doc,
        "3.6 Retrieval Scoring Function",
        "3.6 Retrieval Scoring Function (Layer A MiniLM cosine)",
        "§3.6 heading: Layer A MiniLM cosine, not official BM25",
        changes,
    )

    # Long paragraphs: prepend or replace a sentence only
    prepend_if(
        doc,
        "This chapter presents the formal mathematical basis of the proposed dynamic query routing system:",
        "This chapter first states the mathematics of the historical Layer A router (SVM SHORT/LONG, confidence lights, and MiniLM cosine search). That is not the official frozen retriever. Official M0 scores documents with BM25 after Unicode script routing; official IR metrics are ExactSource Hit@5 and human Success@5 (Section 3.7 and Sections 5.18–5.24).",
        "Ch.3 opening: labeled Layer A; official M0 is BM25, not cosine MiniLM",
        changes,
    )
    prepend_if(
        doc,
        "This chapter describes how ULTRA's θ = 150 character tape is replaced as a whole:",
        "This chapter documents historical Layer A methods (SVM SHORT/LONG, two MiniLM indexes, confidence lights) and the SVM frozen tests (Phases 3A–3B, Phase 2.5, H001–H040 dual-index P@5). Official frozen retrieval is M0 (Unicode detector; URDU/MIXED → Urdu BM25; ROMAN → Method D), evaluated in Sections 5.18–5.24.",
        "Ch.4 opening: Layer A history vs official M0",
        changes,
    )
    prepend_if(
        doc,
        "Restating the gap in prose: first, ULTRA still uses one surface signal and one combined index,",
        "Official IR contribution of this thesis is script-aware BM25 (M0). The two-room SVM is historical Layer A and is not the official retriever.",
        "§2.8: official contribution is M0, not two-room SVM as the IR headline",
        changes,
    )
    replace_if(
        doc,
        "This thesis keeps ULTRA’s two-level idea and replaces the tape with a learned dual-index router.",
        "Historical Layer A keeps ULTRA’s two-level idea and replaces the tape with a learned dual-index SVM. Official M0 instead routes by script to BM25 (Urdu or Method D), as evaluated in Sections 5.18–5.24.",
        "§2.6: thesis does not replace ULTRA only with dual-index SVM; official path is M0 BM25",
        changes,
        whole=False,
    )
    replace_if(
        doc,
        "This thesis addresses that gap directly.",
        "Official IR evaluation addresses the script-routing gap with frozen M0 (Unicode + BM25 + Method D). The SVM-plus-lights design is historical Layer A.",
        "§2.7: official gap-fill is M0, not SVM-plus-lights as the IR system",
        changes,
        whole=False,
    )
    prepend_if(
        doc,
        "After V2 was trained on 409 queries and persisted, Phase 3A verified the eight-feature path",
        "That protocol freezes the SVM classifier, not official M0 BM25. Official M0 retrieval evaluation is Phases 8–12 (Sections 5.18–5.24).",
        "§4.11 body: Phase 3B is SVM freeze, not M0 IR freeze",
        changes,
    )
    prepend_if(
        doc,
        "The retrieval backend uses one news corpus (about 111,860 Urdu articles) and two rooms:",
        "Layer A dual-index (not official M0). Official M0 retrieval is BM25 (Section 4.13 and Sections 5.18–5.24).",
        "§4.2: MiniLM two-room backend labeled Layer A, not official M0",
        changes,
    )
    prepend_if(
        doc,
        "HIGH confidence searches one room. MEDIUM and LOW mix two rooms built with the same encoder",
        "Layer A scoring (not official M0). Official M0 uses BM25, not MiniLM cosine.",
        "§3.6 body: cosine MiniLM labeled Layer A",
        changes,
    )
    prepend_if(
        doc,
        "Because deployment efficiency relative to LLM-based routing is a central claim of this thesis,",
        "The complexity bounds below describe Layer A (SVM + MiniLM/HNSW). Official M0 retrieval cost is BM25 scoring after script routing, not HNSW cosine search.",
        "§3.8: complexity analysis labeled Layer A, not official M0 BM25",
        changes,
    )
    prepend_if(
        doc,
        "Six of eight queries reach 100.00% P@15;",
        "In that same development-only P@15 experiment (not official M0 ExactSource Hit@5 or U Success@5),",
        "§5.2 remaining P@15 paragraph: labeled development-only / not M0",
        changes,
    )
    prepend_if(
        doc,
        "Consolidating the results above, three strengths and three weaknesses of the proposed system",
        "The strengths and weaknesses in this section refer to Layer A (SVM routing and MiniLM pilots), not official M0 ExactSource/Success@5.",
        "§5.10 body: Layer A, not official M0",
        changes,
    )
    prepend_if(
        doc,
        "Beyond the academic contribution of demonstrating that dynamic routing outperforms static thresholding for Urdu IR,",
        "The practitioner implications in this section concern Layer A SVM routing. Official frozen IR deployment is M0 (script-aware BM25), reported in Sections 5.18–5.24.",
        "§5.11: implications labeled Layer A; official IR is M0",
        changes,
    )
    replace_if(
        doc,
        "HIGH / MEDIUM / LOW lights that choose one room, both rooms, or expand-then-both (Section 3.5, 4.6).",
        "Historical Layer A: HIGH / MEDIUM / LOW lights (Sections 3.5, 4.6). Official M0 routing is Unicode script detection, not these lights.",
        "§6.2: lights listed as Layer A, not official M0",
        changes,
    )
    prepend_if(
        doc,
        "Beyond the immediate Urdu news-retrieval setting evaluated in this thesis, the dynamic routing architecture developed here has plausible application",
        "These application remarks concern routing patterns in general. The official frozen retriever for this thesis is M0 (script-aware BM25), not the MiniLM dual-index. The Layer A SVM dual-index is a historical design.",
        "§6.6: official system is M0 BM25, not MiniLM dual-index",
        changes,
    )
    prepend_if(
        doc,
        "Finally, integrating the proposed routing pipeline into a live Urdu search system",
        "If a live system is built from this thesis, the official frozen retriever is M0, not the MiniLM dual-index.",
        "§6.5 closing: live deployment points to M0, not MiniLM",
        changes,
    )

    has_413 = any("4.13 Official M0 evaluation protocol" in (p.text or "") for p in doc.paragraphs)
    if not has_413:
        inserted = False
        for p in doc.paragraphs:
            if (p.text or "").startswith("Phase 2.5 does not train the classifier."):
                h = insert_after(p, "4.13 Official M0 evaluation protocol (Phases 8–12)", "Heading 2")
                insert_after(
                    h,
                    "Official retrieval is frozen M0: URDU/MIXED queries search Urdu BM25; ROMAN queries search Method D. "
                    "The corpus has 111,860 articles. Primary metrics: ExactSource Hit@5 on Phase 2 n=78 and on sealed K001–K040, "
                    "and human Success@5 on sealed U001–U040. Phase 11 did not replace M0. Details and numbers are in Sections 5.18–5.24. "
                    "This protocol is distinct from the Layer A SVM freeze in Section 4.11.",
                    "Normal",
                )
                changes.append("Added §4.13 official M0 evaluation protocol pointer (no new results)")
                inserted = True
                break
        if not inserted:
            changes.append("UNCHANGED: could not insert §4.13")

    n_fig, n_tab = add_seq_captions(doc)
    changes.append(f"Converted {n_fig} figure captions and {n_tab} table captions to Word SEQ fields")
    rebuild_toc(doc, changes)

    try:
        doc.save(str(SRC))
        changes.append(f"Saved {SRC.name}")
    except PermissionError:
        alt = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA_audit.docx"
        doc.save(str(alt))
        changes.append(f"Original locked; saved {alt.name}")
    LOG.write_text("\n".join(f"- {c}" for c in changes) + "\n", encoding="utf-8")
    print("\n".join(changes))


if __name__ == "__main__":
    main()
