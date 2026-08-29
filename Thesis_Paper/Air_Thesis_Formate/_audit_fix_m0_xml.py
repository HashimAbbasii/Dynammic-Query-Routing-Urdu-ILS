# -*- coding: utf-8 -*-
"""Wipe leftover TOC field XML from moved M0 paragraphs. Presentation only."""
from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
SRC = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"
LOG = HERE / "_audit_changes.txt"

HEADING_MAP = {
    "5.18 Official frozen retrieval system (M0)": "5.18 Official frozen retrieval system (M0)",
    "5.19 Phase 2 development/validation known-item (n = 78)": "5.19 Phase 2 development/validation known-item (n = 78)",
    "5.20 Phase 11 ablation: M0 remains official": "5.20 Phase 11 ablation: M0 remains official",
    "5.21 Phase 12 new known-item evaluation (K001–K040)": "5.21 Phase 12 new known-item evaluation (K001–K040)",
    "5.22 Phase 12 naturalistic human evaluation (U001–U040)": "5.22 Phase 12 naturalistic human evaluation (U001–U040)",
    "5.23 Final comparison of evaluation settings": "5.23 Final comparison of evaluation settings",
    "5.24 Discussion of the official IR results": "5.24 Discussion of the official IR results",
}

TABLE_RENAMES = (
    ("Table 3. Phase 2 n=78", "Table 22. Phase 2 n=78"),
    ("Table 4. Phase 11 M0–M4", "Table 23. Phase 11 M0–M4"),
    ("Table 5. Phase 12 K ExactSource", "Table 24. Phase 12 K ExactSource"),
    ("Table 6. Phase 12 U human", "Table 25. Phase 12 U human"),
    ("Table 7. Script-wise U Success@5", "Table 26. Script-wise U Success@5"),
    ("Table 8. Do not average rows.", "Table 27. Do not average rows."),
)


def try_style(p, name: str) -> None:
    try:
        p.style = name
    except Exception:
        pass


def wipe_and_set(p, text: str) -> None:
    el = p._p
    p_pr = el.find(qn("w:pPr"))
    for child in list(el):
        if child is not p_pr:
            el.remove(child)
    p.add_run(text)


def copy_ppr(src, dest) -> None:
    src_pr = src._p.find(qn("w:pPr"))
    dest_el = dest._p
    old = dest_el.find(qn("w:pPr"))
    if old is not None:
        dest_el.remove(old)
    if src_pr is not None:
        dest_el.insert(0, deepcopy(src_pr))


def payload(t: str) -> str:
    t = t.replace("\u00a0", " ")
    for old, new in TABLE_RENAMES:
        t = t.replace(old, new)
    m = re.search(r"400 judgments\)\s*\t?\s*76(.*)$", t, re.S)
    if m:
        rest = m.group(1)
        if rest.startswith("."):
            rest = "5" + rest
        t = rest.strip()
    for key, clean in HEADING_MAP.items():
        if key in t and not t.startswith("Table"):
            return clean
        mashed = key[1:]  # ".18 Official..."
        if t.startswith(mashed):
            return clean
    return t.strip()


def classify(text: str) -> str:
    if text.startswith("5.1") and text[:4] in ("5.18", "5.19", "5.20", "5.21", "5.22", "5.23", "5.24"):
        return "Heading 2"
    if text.startswith("Table 2") and len(text) > 8 and text[6:8].isdigit():
        return "Caption"
    return "Normal"


def has_field(p) -> bool:
    return p._p.find(".//" + qn("w:fldChar")) is not None or p._p.find(qn("w:fldChar")) is not None


def add_seq_if_plain(p) -> bool:
    raw = (p.text or "").strip()
    m = re.match(r"^(Figure|Table)\s+[\d.]+\.\s*(.*)$", raw, re.S)
    if not m:
        return False
    if p._p.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar") is not None:
        return False
    kind, title = m.group(1), m.group(2).strip()
    from docx.oxml import OxmlElement

    wipe_and_set(p, "")
    p.add_run(f"{kind} ")
    run_begin = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(begin)
    run_instr = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {kind} \\* ARABIC "
    run_instr._r.append(instr)
    run_sep = p.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(sep)
    p.add_run("1")
    run_end = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run_end._r.append(end)
    p.add_run(f". {title}")
    return True


def main() -> None:
    doc = Document(str(SRC))
    h2_src = cap_src = norm_src = None
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        t = (p.text or "").strip()
        if h2_src is None and st.startswith("Heading") and t.startswith("5.17 Frozen"):
            h2_src = p
        if cap_src is None and st == "Caption":
            cap_src = p
        if norm_src is None and st == "Normal" and len(t) > 40 and "Layer A" in t:
            norm_src = p

    start = end = None
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        t = p.text or ""
        if "Official frozen retrieval system (M0)" in t and start is None:
            start = i
        if start is not None and "Do not retune M0 on K, U, or H001–H040" in t:
            end = i
            break
    if start is None or end is None:
        raise SystemExit("M0 block not found")

    n = 0
    for i in range(start, end + 1):
        p = paras[i]
        text = payload(p.text or "")
        kind = classify(text)
        if kind == "Heading 2" and h2_src is not None:
            copy_ppr(h2_src, p)
        elif kind == "Caption" and cap_src is not None:
            copy_ppr(cap_src, p)
        elif kind == "Normal" and norm_src is not None:
            copy_ppr(norm_src, p)
        wipe_and_set(p, text)
        try_style(p, kind)
        if kind == "Caption":
            add_seq_if_plain(p)
        n += 1

    extra = []
    extra.append("Wiped TOC hyperlink/field leftovers from %d M0 paragraphs (%d-%d)" % (n, start, end))
    extra.append("Restored clean Heading 2 / Caption / Normal styles for §§5.18–5.24")

    # Re-verify order
    idx_517 = idx_518 = idx_ch6 = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        st = p.style.name if p.style else ""
        if t.startswith("5.17 Frozen") and st.startswith("Heading") and idx_517 is None:
            idx_517 = i
        if t.startswith("5.18 Official") and st.startswith("Heading"):
            idx_518 = i
        if t == "Chapter 6: Conclusions and Recommendations" and st.startswith("Heading"):
            idx_ch6 = i
    extra.append("ORDER 5.17=%s 5.18=%s Ch6=%s OK=%s" % (idx_517, idx_518, idx_ch6, idx_517 is not None and idx_518 is not None and idx_ch6 is not None and idx_517 < idx_518 < idx_ch6))

    mash = 0
    for p in doc.paragraphs:
        t = p.text or ""
        if "5.17 Frozen held-out dual-index P@5 (400 judgments)" in t and ("5.18" in t or "76.18" in t or "Table 22" in t):
            mash += 1
    extra.append("TOC_MASH_COUNT=%d" % mash)

    try:
        doc.save(str(SRC))
        extra.append("Saved Hashim_Shazad_243259_AU_Thesis_ULTRA.docx")
    except PermissionError:
        alt = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA_audit.docx"
        doc.save(str(alt))
        extra.append("Original locked; saved " + alt.name)

    prev = LOG.read_text(encoding="utf-8") if LOG.exists() else ""
    LOG.write_text(prev + "\n".join("- " + e for e in extra) + "\n", encoding="utf-8")
    Path(HERE / "_audit_verify2.txt").write_text("\n".join(extra), encoding="utf-8")
    print("ok")


if __name__ == "__main__":
    main()
