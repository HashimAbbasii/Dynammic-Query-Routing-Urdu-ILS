# -*- coding: utf-8 -*-
"""Restore AU logo, drop duplicate cover, keep supervisor on the title page, fix caption numbers."""
from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips
from docx.text.paragraph import Paragraph

HERE = Path(__file__).resolve().parent
SRC = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"
BAK = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.pre_finalize.bak.docx"
LOGO_BAK_DOC = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.pre_onestory.bak.docx"
LOGO_PNG = HERE / "_air_logo_from_backup.png"


def has_drawing(p: Paragraph) -> bool:
    return bool(p._p.findall(".//" + qn("w:drawing")))


def has_page_break(p: Paragraph) -> bool:
    xml = p._p.xml
    return 'w:type="page"' in xml


def has_sectpr(p: Paragraph) -> bool:
    return p._p.find(qn("w:sectPr")) is not None


def delete_paragraph(p: Paragraph) -> None:
    el = p._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def set_text(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def compact(p: Paragraph, before_pt: float = 0, after_pt: float = 4, keep_next: bool = True) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.keep_together = True
    pf.keep_with_next = keep_next
    pf.page_break_before = False


def restore_logo(docx_path: Path, logo: Path) -> None:
    tmp = docx_path.with_name(docx_path.stem + "._logo.tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        data = logo.read_bytes()
        for name in zin.namelist():
            zout.writestr(name, data if name == "word/media/image1.png" else zin.read(name))
    shutil.copy2(tmp, docx_path)
    tmp.unlink(missing_ok=True)


def extract_logo() -> None:
    if LOGO_PNG.exists() and LOGO_PNG.stat().st_size > 10000:
        return
    if not LOGO_BAK_DOC.exists():
        raise SystemExit("no backup to restore Air University logo from")
    with zipfile.ZipFile(LOGO_BAK_DOC) as z:
        LOGO_PNG.write_bytes(z.read("word/media/image1.png"))


def drop_duplicate_cover(doc: Document) -> int:
    paras = list(doc.paragraphs)
    first_img = next((i for i, p in enumerate(paras) if has_drawing(p)), None)
    if first_img is None:
        raise SystemExit("no cover logo paragraph")
    first_break = next((i for i, p in enumerate(paras) if i > first_img and has_page_break(p)), None)
    if first_break is None:
        raise SystemExit("no page break after first cover")
    second_img = next((i for i, p in enumerate(paras) if i > first_break and has_drawing(p)), None)
    if second_img is None:
        print("no second cover image; skip duplicate delete")
        return 0
    # Delete first cover through the blanks before the second logo. Never touch sectPr.
    removed = 0
    for p in paras[first_img:second_img]:
        if has_sectpr(p):
            continue
        delete_paragraph(p)
        removed += 1
    print(f"removed {removed} paragraphs from duplicate cover")
    return removed


def polish_title_page(doc: Document) -> None:
    # Remaining first drawing is the real cover logo.
    paras = list(doc.paragraphs)
    logo_i = next(i for i, p in enumerate(paras) if has_drawing(p))
    block = []
    for p in paras[logo_i:]:
        if has_sectpr(p) or (p.text or "").strip() == "CERTIFICATE":
            break
        block.append(p)

    for p in block:
        t = (p.text or "").strip()
        if has_drawing(p):
            compact(p, 0, 6, True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if not t:
            compact(p, 0, 0, True)
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if t.startswith("ADAPTIVE"):
            compact(p, 6, 8, True)
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(16)
        elif t == "by":
            compact(p, 0, 2, True)
        elif t == "Hashim Shazad":
            compact(p, 0, 6, True)
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(16)
        elif t.startswith("A thesis submitted"):
            compact(p, 0, 2, True)
        elif t.startswith("Master of Science"):
            compact(p, 0, 8, True)
            for r in p.runs:
                r.bold = True
        elif t == "SUPERVISOR":
            compact(p, 8, 2, True)
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(12)
        elif t.startswith("Dr. Adnan"):
            compact(p, 0, 8, True)
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(14)
        elif t == "AIR UNIVERSITY":
            compact(p, 4, 0, True)
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(14)
        elif t == "ISLAMABAD":
            compact(p, 0, 2, True)
            for r in p.runs:
                r.bold = True
        elif t.startswith("©"):
            set_text(p, "© Hashim Shazad, 2026. All rights reserved.")
            compact(p, 8, 0, False)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            compact(p, 0, 4, True)


def fix_seq_numbers(doc: Document) -> dict[str, int]:
    counters = {"Figure": 0, "Table": 0, "Algorithm": 0}
    instr = None
    after_sep = False
    for p in doc.paragraphs:
        for el in p._p.iter():
            if el.tag == qn("w:instrText") and el.text and "SEQ" in el.text:
                m = re.search(r"SEQ\s+(Figure|Table|Algorithm)", el.text)
                if m:
                    instr = m.group(1)
                    after_sep = False
            elif el.tag == qn("w:fldChar") and el.get(qn("w:fldCharType")) == "separate":
                after_sep = True
            elif after_sep and instr and el.tag == qn("w:t"):
                counters[instr] += 1
                el.text = str(counters[instr])
                after_sep = False
                instr = None
    return counters


def keep_headings_with_body(doc: Document) -> None:
    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        if name in {"Heading 1", "Heading 2", "Heading 3", "Caption"}:
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.widow_control = True


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"missing {SRC}")
    shutil.copy2(SRC, BAK)
    extract_logo()

    doc = Document(str(SRC))
    drop_duplicate_cover(doc)
    polish_title_page(doc)
    counts = fix_seq_numbers(doc)
    print("caption SEQ cache", counts)
    keep_headings_with_body(doc)

    out = SRC
    try:
        doc.save(str(out))
    except PermissionError:
        out = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA_FINAL.docx"
        doc.save(str(out))
        print("original locked; wrote", out)

    restore_logo(out, LOGO_PNG)
    print("restored Air University logo into", out)
    print("backup", BAK)


if __name__ == "__main__":
    main()
