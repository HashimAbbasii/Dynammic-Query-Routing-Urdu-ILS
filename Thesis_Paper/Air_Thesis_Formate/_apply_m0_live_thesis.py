# -*- coding: utf-8 -*-
"""Align the LIVE AU Word thesis to frozen M0 IR results. No experiments. No Git."""
from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

HERE = Path(__file__).resolve().parent
SRC = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.docx"
BAK = HERE / "Hashim_Shazad_243259_AU_Thesis_ULTRA.pre_m0_cleanup.bak.docx"


def set_text(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def insert_after(paragraph, text: str, style_name: str):
    new_el = deepcopy(paragraph._p)
    paragraph._p.addnext(new_el)
    np = Paragraph(new_el, paragraph._parent)
    try:
        np.style = style_name
    except Exception:
        pass
    set_text(np, text)
    return np


def replace_prefixed(doc, jobs):
    hit, miss = [], []
    used = set()
    for prefix, new in jobs:
        found = False
        for i, p in enumerate(doc.paragraphs):
            if i in used:
                continue
            t = p.text or ""
            if t.startswith(prefix) or (len(prefix) > 48 and prefix in t[:200]):
                set_text(p, new)
                used.add(i)
                hit.append(prefix[:70])
                found = True
                break
        if not found:
            miss.append(prefix[:80])
    return hit, miss


ABSTRACT = (
    "Urdu information retrieval is difficult because users may type native Perso-Arabic script, "
    "informal Roman Urdu, or mixed forms, while many systems still apply one retrieval path to every query. "
    "This thesis proposes adaptive dynamic query routing for the ULTRA news-search setting. The official frozen "
    "system, M0, uses a Unicode script detector over a corpus of 111,860 cleaned Urdu news articles: URDU and MIXED "
    "queries search an Urdu BM25 index, and ROMAN queries search a Method D romanized-document BM25 index. "
    "Query-side variants M1–M4 did not improve the primary known-item score, so M0 was not replaced. "
    "On the Phase 2 development/validation known-item set, ExactSource Hit@5 is 68/78 = 87.18%. That figure is genuine "
    "for title-derived known-item search on that pool; it is not unseen usefulness and not overall system accuracy. "
    "Independent Phase 12 evaluation of the same freeze produced two further, separate results: ExactSource Hit@5 = "
    "27/40 = 67.50% on new known-item queries (K001–K040), and human Success@5 = 23/40 = 57.50% on naturalistic queries "
    "(U001–U040; conservative P@5 = 0.2050, nDCG@5 = 0.6460, MRR = 0.4542). Success@5 is human relevance, not ExactSource "
    "Hit@5. These three percentages must not be averaged. In the U sample, Urdu-script queries succeeded in 17/18 cases "
    "versus 6/18 Roman and 0/4 mixed. The contribution is a frozen, reproducible script-aware retriever with strong "
    "controlled known-item performance and an explicitly measured Roman Urdu limitation."
)

CH5_OPEN = (
    "This chapter contains two layers that must not be mixed. Sections 5.1–5.17 report historical SVM SHORT/LONG "
    "routing and MiniLM dual-index pilots (Layer A). Those numbers are classification accuracy, development P@15, or "
    "dense-index P@5. They are not official M0 ExactSource Hit@5 and not U Success@5. Sections 5.18–5.24 report the "
    "official frozen retrieval system M0: Unicode routing, Urdu BM25, and Method D. Official IR headlines are "
    "68/78 = 87.18% ExactSource Hit@5 (development/validation known-item), 27/40 = 67.50% ExactSource Hit@5 (new K), "
    "and 23/40 = 57.50% human Success@5 (new U). Do not average those three percentages. Do not treat 87.18% as unseen "
    "or human usefulness. Do not treat development 90.00% P@15 as frozen M0 retrieval performance."
)

P15_NOTE = (
    "DEVELOPMENT-ONLY (not official M0). Because the original ULTRA framework has no Roman Urdu handling, it scores "
    "0.00% P@15 on Roman Urdu queries by construction, pulling its overall weighted P@15 down to 43.75% in this small "
    "development retrieval experiment. The proposed development pipeline matches ULTRA on native Urdu (87.50%) while "
    "adding 92.50% P@15 on eight Roman Urdu queries, raising overall P@15 to 90.00% in that experiment only. This is "
    "not ExactSource Hit@5, not Phase 12 Success@5, and not the official frozen M0 evaluation."
)

DISC_512 = (
    "Taken together, Layer A shows that θ = 150 is the wrong length switch and that an SVM can learn SHORT versus LONG "
    "on development data, with only a small frozen classification edge versus word count (86% vs 84%) and no dual-index "
    "P@5 gain on H001–H040. That story is not the official retrieval evaluation. Official M0 results are in Sections "
    "5.18–5.24: 87.18% ExactSource Hit@5 on the development/validation known-item pool, 67.50% ExactSource Hit@5 on new "
    "known-item queries, and 57.50% human Success@5 on naturalistic queries. Roman Urdu is the main usefulness gap "
    "(U: 6/18 versus Urdu 17/18). Development 92.50% P@15 and 90.00% overall P@15 remain supporting Layer A figures, "
    "not the M0 headline."
)

CONC_1 = (
    "This thesis evaluated Urdu news retrieval under a frozen lexical system (M0) after earlier work on learned "
    "SHORT/LONG routing. Official M0 uses Unicode script detection, Urdu BM25 for URDU/MIXED queries, and Method D "
    "for ROMAN queries over 111,860 articles. M0 was not changed after Phase 12. M1–M4 did not improve n=78 Hit@5."
)
CONC_2 = (
    "The frozen M0 system achieved strong performance on the development/validation known-item evaluation "
    "(68/78 = 87.18% ExactSource Hit@5). However, evaluation on newly constructed queries produced lower performance: "
    "67.50% ExactSource Hit@5 (27/40) on new known-item queries and 57.50% human Success@5 (23/40) on naturalistic "
    "queries. This indicates that the development/validation score should not be interpreted as a direct estimate of "
    "naturalistic user-level retrieval usefulness. ExactSource Hit@5 and Success@5 measure different properties and "
    "must not be averaged with each other or with 87.18%."
)
CONC_3 = (
    "ULTRA therefore demonstrates effective script-aware retrieval on the development/validation known-item benchmark, "
    "while new unseen evaluations reveal a substantial generalization gap and particularly weak Roman-script performance "
    "(U Success@5: Urdu 17/18, Roman 6/18, Mixed 0/4). The contribution is not 87% real-world accuracy. H001–H040 "
    "Success@5 = 62.5% is diagnostic only."
)

LIM_INTRO = (
    "Limitations of the official M0 evaluation are stated first. Historical SVM-layer limits (Phase 3B power, HIGH-tier "
    "demonstration) remain valid for Layer A but do not replace the IR limits below."
)
LIM_DATA = (
    "Unseen evaluation size. New known-item (K) and naturalistic (U) sets each have n = 40. Point estimates have wide "
    "uncertainty. Numerators and denominators must be reported. This evaluation should not be interpreted as universal "
    "real-world accuracy."
)
LIM_METRICS = (
    "Known-item versus naturalistic evaluation. ExactSource Hit@5 requires a pre-assigned source document. Human "
    "Success@5 asks whether any Top-5 document is relevant or partially relevant. The drop from 87.18% to 57.50% is not "
    "one metric getting worse. Development/validation and unseen results differ for both scientific and sampling reasons."
)
LIM_ROMAN = (
    "Roman Urdu performance is substantially weaker than Urdu-script performance on the sealed U sample (6/18 versus "
    "17/18). Ordinary and chat spelling diverge from Method D’s title_roman construction. The mixed-script subset was "
    "very small (n = 4, all unsuccessful) and cannot support a population rate."
)
LIM_HUMAN = (
    "Naturalistic queries require human relevance judgments. U labels are from one annotator; there is no inter-annotator "
    "agreement. nDCG@5 with topical C-gain = 1 can look high when Success@5 fails and is not the usefulness headline."
)

M0_BLOCKS = [
    ("Heading 2", "5.18 Official frozen retrieval system (M0)"),
    (
        "Normal",
        "The official retriever is M0. Queries are not rewritten. A Unicode letter-count detector labels URDU, ROMAN, "
        "or MIXED. URDU and MIXED search Urdu BM25 over article text. ROMAN searches Method D BM25 over romanized "
        "documents. BM25 uses k1 = 1.5 and b = 0.75. The corpus contains 111,860 news articles. The Roman dictionary "
        "has 198 keys. Phase 11 query-side models M1–M4 did not improve n=78 ExactSource Hit@5; M0 was not replaced.",
    ),
    (
        "Normal",
        "ExactSource Hit@5 asks whether the exact source article appears in the Top-5 and requires a source_doc_id. "
        "Human Success@5 asks whether any Top-5 article is labelled A (relevant) or B (partially relevant). These "
        "metrics are not interchangeable and are not averaged.",
    ),
    ("Heading 2", "5.19 Phase 2 development/validation known-item (n = 78)"),
    (
        "Normal",
        "Table 3. Phase 2 n=78 ExactSource results (official M0). Evaluation type: known-item, development/validation. "
        "M0 ExactSource Hit@5 = 68/78 = 87.18%. Urdu-only BM25 (no Roman path) = 0.5897. Roman subset Method A = 0/23. "
        "Roman subset Method D = 22/23.",
    ),
    (
        "Normal",
        "The frozen ULTRA system achieved an ExactSource Hit@5 of 87.18% (68/78) on the Phase 2 development/validation "
        "known-item evaluation set. This result is genuine within that protocol: the designated source was in the Top-5 "
        "for 68 of 78 title-derived queries. It is not real-world accuracy, not human usefulness, and not unseen "
        "natural-query performance. Roman QTRN strings are Phase 2 title_roman, not chat-style Roman Urdu.",
    ),
    ("Heading 2", "5.20 Phase 11 ablation: M0 remains official"),
    (
        "Normal",
        "Table 4. Phase 11 M0–M4 ablation on the same n=78 known-item pool. M0, M1, M2, M3, and M4 all scored "
        "68/78 = 87.18% ExactSource Hit@5. Roman-train Hit@5 remained 61/64 = 95.31%. M1–M4 nDCG@5 was slightly below "
        "M0. M1 is a gate-passing candidate, not an improvement, and is not the official system.",
    ),
    ("Heading 2", "5.21 Phase 12 new known-item evaluation (K001–K040)"),
    (
        "Normal",
        "Table 5. Phase 12 K ExactSource results, frozen M0, sealed before retrieval. ExactSource Hit@1 = 20/40 = 50.00%; "
        "Hit@5 = 27/40 = 67.50% (primary); Hit@10 = 28/40 = 70.00%; Hit@50 = 30/40 = 75.00%. This is a new known-item "
        "evaluation. It does not replace 68/78 and is not human Success@5.",
    ),
    (
        "Normal",
        "Descriptive detector split, not used for tuning: Urdu-script K titles 26/28; ordinary Roman titles 1/12. The "
        "drop from 87.18% to 67.50% is concentrated on ordinary Roman title queries, not on native-script Urdu BM25.",
    ),
    ("Heading 2", "5.22 Phase 12 naturalistic human evaluation (U001–U040)"),
    (
        "Normal",
        "Table 6. Phase 12 U human results, frozen M0 Top-5. Primary usefulness metric: Success@5 = 23/40 = 57.50%. "
        "Conservative P@5 = 0.2050. nDCG@5 = 0.6460 (gains A=3, B=2, C=1, D=E=0; secondary). MRR = 0.4542. This is "
        "human relevance, not ExactSource Hit@5. nDCG@5 is not the usefulness headline because C still has gain 1.",
    ),
    (
        "Normal",
        "Table 7. Script-wise U Success@5 (descriptive only). URDU 17/18 = 94.44%; ROMAN 6/18 = 33.33%; MIXED 0/4 = 0%. "
        "These findings describe this sealed sample. They do not licence retuning Method D on U failures, and they do "
        "not prove that Method D is universally bad. They do show that Urdu-script needs were usually met and Roman "
        "and mixed needs often were not. Mixed n=4 is too small for a population rate.",
    ),
    (
        "Normal",
        "H001–H040 human Success@5 = 25/40 = 62.5% is a diagnostic on earlier trap queries. ExactSource Hit@5 on that "
        "set is undefined (no source_doc_id). It is not the official unseen usefulness result and must not be combined "
        "with U.",
    ),
    ("Heading 2", "5.23 Final comparison of evaluation settings"),
    (
        "Normal",
        "Table 8. Do not average rows. Phase 2 n=78 known-item ExactSource Hit@5 = 68/78 = 87.18% (development/validation). "
        "Phase 12 K n=40 known-item ExactSource Hit@5 = 27/40 = 67.50% (new). Phase 12 U n=40 human Success@5 = "
        "23/40 = 57.50% (new naturalistic). H001–H040 Success@5 = 25/40 = 62.5% (diagnostic only).",
    ),
    ("Heading 2", "5.24 Discussion of the official IR results"),
    (
        "Normal",
        "The frozen M0 system achieved strong performance on the development/validation known-item evaluation "
        "(68/78, 87.18% ExactSource Hit@5). However, evaluation on newly constructed queries produced lower performance: "
        "67.50% ExactSource Hit@5 on new known-item queries and 57.50% human Success@5 on naturalistic queries. This "
        "indicates that the development/validation score should not be interpreted as a direct estimate of naturalistic "
        "user-level retrieval usefulness.",
    ),
    (
        "Normal",
        "These results suggest a staircase of task difficulty and query-form mismatch, not a single accuracy that collapsed. "
        "87.18% asks whether a title-derived known-item, including title_roman, recovered its source on the freeze pool. "
        "67.50% asks the same known-item question on independently sampled titles; Urdu titles remain high and ordinary "
        "Roman titles do not. 57.50% asks a different question: on natural needs with no gold article, was anything in "
        "the Top-5 useful? Conservative P@5 = 0.2050 suggests many successes are a single useful document. Query mix "
        "matters: a Roman-heavy sample will look worse than an Urdu-only test even if the Urdu component is strong.",
    ),
    (
        "Normal",
        "In official M0, adaptive dynamic query routing is script-conditional index selection, not the SVM SHORT/LONG "
        "switch. Method D is necessary for title_roman-like Roman known-item search (22/23 on the development Roman "
        "subset versus 0/23 for Method A) and is insufficient for ordinary and chat Roman Urdu on new K and U. Phase 11 "
        "showed that small allowed query expansions did not move 68/78. These results suggest that Roman Urdu remains "
        "the major limitation of the frozen system. Do not hide 57.50%. Do not retune M0 on K, U, or H001–H040.",
    ),
]


def insert_blocks_before_chapter6(doc):
    ch6 = None
    for p in doc.paragraphs:
        if (p.text or "").startswith("Chapter 6:"):
            ch6 = p
            break
    if ch6 is None:
        raise SystemExit("Chapter 6 heading not found")
    # Walk backward: insert after the paragraph immediately preceding Chapter 6
    prev = ch6._p.getprevious()
    from docx.oxml.ns import qn

    # Find preceding w:p
    while prev is not None and prev.tag != qn("w:p"):
        prev = prev.getprevious()
    if prev is None:
        raise SystemExit("no paragraph before Chapter 6")
    anchor = Paragraph(prev, ch6._parent)
    cur = anchor
    for style, text in M0_BLOCKS:
        cur = insert_after(cur, text, style)
    return len(M0_BLOCKS)


def main():
    if not SRC.exists():
        raise SystemExit(f"missing {SRC}")
    if not BAK.exists():
        shutil.copy2(SRC, BAK)
        print("safety copy", BAK.name)
    else:
        print("safety copy already exists", BAK.name)

    doc = Document(str(SRC))
    jobs = [
        (
            "Urdu information retrieval (IR) systems face persistent challenges arising from script variability",
            ABSTRACT,
        ),
        (
            "This thesis extends ULTRA by replacing the θ = 150 character tape as a whole",
            "This thesis extends ULTRA by treating script as the official routing decision for lexical news search. "
            "The frozen system M0 sends URDU and MIXED queries to Urdu BM25 and ROMAN queries to Method D. Earlier SVM "
            "SHORT/LONG work (two semantic rooms and confidence lights) is retained as historical Layer A. Official IR "
            "evaluation uses ExactSource Hit@5 and human Success@5, not development P@15 as the headline.",
        ),
        (
            "Existing Urdu IR systems, including ULTRA, still treat routing as a length tape.",
            "The practical problem is bilingual and script-mixed Urdu news search over 111,860 articles. Users issue "
            "native-script Urdu, Roman Urdu with non-standard spelling, and occasional mixed-script queries. A single "
            "Urdu-text BM25 index is a poor match for Roman queries (development Roman known-item Hit@5 was 0/23 without "
            "a Roman path). A static length threshold does not decide which script index to open. Classification accuracy "
            "on SHORT versus LONG is not the same as finding a useful news article.",
        ),
        (
            "Those failures show up in three concrete ways.",
            "The scientific problem is therefore to freeze a retrieval system, report known-item recovery under a defined "
            "protocol, then report generalization on new known-item queries and naturalistic human usefulness without "
            "mixing those metrics. Chapter 4 specifies M0. Chapter 5 reports 68/78, 27/40, and 23/40 separately.",
        ),
        (
            "A search system that cannot decide which room to open will under-serve users even if the encoder is strong.",
            "A search system that cannot decide which script index to open will miss Roman Urdu even when the article exists. "
            "The motivation is a frozen, inspectable BM25 pipeline with honest evaluation: development known-item, new "
            "known-item, and human usefulness kept apart. Overstating 87.18% as real-world accuracy would hide the Phase 12 gap.",
        ),
        (
            "The gap is joint, not three separate papers.",
            "The evaluation gap this thesis addresses is a frozen Urdu/Roman lexical router with three non-averaged layers: "
            "development/validation known-item ExactSource Hit@5, new known-item ExactSource Hit@5, and naturalistic human "
            "Success@5. English query routing and Urdu collections exist; this freeze protocol did not.",
        ),
        (
            "RQ1: Can a small set of query features",
            "RQ1: Can a script-aware lexical pipeline (Urdu BM25 + Method D) recover known news articles from title-derived "
            "queries on a development/validation pool?",
        ),
        (
            "RQ2: Does that learned router beat",
            "RQ2: Does that known-item score transfer to a new sealed known-item sample (K001–K040) written independently "
            "of the freeze set?",
        ),
        (
            "RQ3: Can calibrated confidence act as traffic lights",
            "RQ3: How often is frozen M0 useful (at least one A or B in the Top-5) on new naturalistic queries with no gold "
            "article (U001–U040)? RQ4: Do query-side Roman expansions M1–M4 improve n=78 ExactSource Hit@5 enough to replace "
            "M0? Historical Layer A (SVM SHORT/LONG and MiniLM P@5) is reported separately and is not the official IR headline.",
        ),
        (
            "Dynamic query classification: Train an SVM that labels SHORT",
            "Freeze an official retrieval system (M0) with documented hashes, Unicode routing, Urdu BM25, and Method D.",
        ),
        (
            "Roman Urdu support: Detect Latin-script Urdu and transliterate",
            "Report ExactSource Hit@5 on Phase 2 development/validation known-item queries (n=78) and on sealed K001–K040, "
            "and human Success@5 on sealed U001–U040, without using those queries for tuning.",
        ),
        (
            "Two rooms and lights: Search a headline index or a full-article index",
            "Report the Phase 11 M0–M4 ablation and keep M0 official if Hit@5 does not improve.",
        ),
        (
            "Comparative evaluation: Keep Phase 3B frozen (86% vs 84%).",
            "State limitations, especially Roman and mixed-script performance, without hiding weaker Phase 12 numbers.",
        ),
        (
            "A dual-index SVM router: SHORT opens headlines, LONG opens full articles.",
            "A frozen script-aware BM25 architecture (M0) with Unicode routing and Method D. Official IR: ExactSource Hit@5 "
            "= 68/78 = 87.18% (development/validation known-item); 27/40 = 67.50% (new K); human Success@5 = 23/40 = 57.50% "
            "(new U). M1–M4 did not improve 68/78.",
        ),
        (
            "An expanded Roman Urdu transliteration dictionary (198 on-disk pairs; earlier drafts cited 179), improving Roman Urdu retrieval coverage in the development P@15 experiment",
            "An expanded Roman Urdu dictionary (198 on-disk pairs). Development P@15 of 92.5% on eight Roman queries is a "
            "historical supporting experiment, not official M0 Hit@5 or Success@5.",
        ),
        (
            "This thesis is scoped to query routing and to making that route change retrieval.",
            "Official retrieval evaluation is news-domain M0 over 111,860 articles. It does not replace M0 with M1, does not "
            "treat H001–H040 as the primary unseen test, and does not claim legal, medical, or live web search. Historical "
            "SVM dual-index work remains in Sections 5.1–5.17 as Layer A only.",
        ),
        (
            "Chapter 5 — Results and Discussion. Keeps development 100% CV separate from frozen 86/84, then reports 60/20 traps and dual-index P@5.",
            "Chapter 5 — Results and Discussion. Labels Layer A (SVM/P@15/MiniLM) as historical, then reports official M0: "
            "87.18% ExactSource Hit@5 (n=78), 67.50% ExactSource Hit@5 (K), and 57.50% Success@5 (U).",
        ),
        (
            "Chapter 4 — Methods Developed. Describes two indexes, the SVM router, lights, 409-query V2 training, trap-augmented twelve-feature training, and frozen evaluation protocols.",
            "Chapter 4 — Methods Developed. Describes historical SVM routing and, for official IR, frozen M0 script routing, "
            "Method D, and the Phase 8–12 evaluation protocol.",
        ),
        (
            "Chapter 3 — Mathematical Formulation. Presents the V2 feature vector, the SVM and lights, and cosine scoring in each room (headline matrix and full-article index).",
            "Chapter 3 — Mathematical Formulation. Presents historical SVM mathematics (Layer A) and the official M0 metrics: "
            "BM25, ExactSource Hit@k, and human Success@5.",
        ),
        (
            "This chapter keeps three evaluation layers apart. Sections 5.1–5.12 are development:",
            CH5_OPEN,
        ),
        (
            "Because the original ULTRA framework has no Roman Urdu handling, it scores 0.00% P@15 on Roman Urdu queries by construction",
            P15_NOTE,
        ),
        (
            "Relative to the base ULTRA framework (Bashir et al., 2026) that this thesis directly extends, the reported native-Urdu P@15 (87.50%) is unchanged",
            "Relative to the base ULTRA framework (Bashir et al., 2026), the development P@15 comparison (43.75% to 90.00% overall in a small experiment) is Layer A only. Official M0 comparison on n=78 is ExactSource Hit@5 = 68/78 versus Urdu-only BM25 0.5897, not that P@15 table.",
        ),
        (
            "Taken together, the honest story is this. θ = 150 never fires LONG",
            DISC_512,
        ),
        (
            "Third, the Roman Urdu transliteration module extends retrieval usability to code-mixed and transliterated queries (92.50% P@15, raising overall system P@15 from 43.75% to 90.00%)",
            "Third, development P@15 of 92.50% on eight Roman queries (overall 90.00% in that experiment) is not official M0 performance. Official Roman evidence is Method D 22/23 on development title_roman known-items, K Roman ExactSource 1/12, and U Roman Success@5 6/18.",
        ),
        (
            "The frozen primary evaluation (n = 50) is the methodologically strongest routing result in this thesis.",
            "The frozen primary SVM evaluation (n = 50) is the methodologically strongest Layer A classification result. It is not official M0 ExactSource Hit@5. Official IR generalization is Phase 12 K and U (Sections 5.21–5.22).",
        ),
        (
            "As a diagnostic robustness check beyond the training split, a 50-query native-Urdu batch was evaluated during development (Figure 5.12). The deployed V2 model records 98.00% accuracy on this diagnostic set",
            "As a diagnostic robustness check for the SVM router, a 50-query native-Urdu batch was evaluated during development (Figure 5.12). V2 records 98.00% routing accuracy on that diagnostic set. This is not official M0 Hit@5 and is not unseen human usefulness. The official IR tests are n=78 ExactSource, K, and U.",
        ),
        (
            "This thesis addressed ULTRA as a whole: not only the θ = 150 decision, but the search that must follow it.",
            CONC_1,
        ),
        (
            "Headline numbers: frozen Phase 3B 86.00% (V2 SVM) vs 84.00% (six-word rule)",
            CONC_2,
        ),
        (
            "RQ1 is supported with a limit: script and length features plus why/how/fact flags predict room choice when those flags fire.",
            CONC_3,
        ),
        (
            "Two-room SVM routing, frozen 86/84 (V2) and frozen 60/20 (traps).",
            "Official M0: 68/78 ExactSource Hit@5 (development/validation), 27/40 ExactSource Hit@5 (K), 23/40 Success@5 (U). Historical SVM 86/84 and trap 60/20 remain Layer A (Sections 5.13 and 5.16).",
        ),
        (
            "An expanded Roman Urdu dictionary (198 on-disk pairs), with development Roman P@15 of 92.5% as a supporting result (Section 5.2).",
            "An expanded Roman Urdu dictionary (198 on-disk pairs). Development Roman P@15 of 92.5% is supporting Layer A only (Section 5.2).",
        ),
        (
            "Frozen Phase 3B (50 primary) plus a frozen 40-query trap set and 400 dual-index relevance rows (Sections 5.16–5.17).",
            "Official Phase 12 K and U (Sections 5.21–5.22), plus diagnostic H001–H040 Success@5 = 62.5%. Historical dual-index P@5 remains Section 5.17 and is not M0 Success@5.",
        ),
        (
            "Several limitations qualify these findings and motivate the future work outlined below.",
            LIM_INTRO,
        ),
        (
            "Dataset scale and topical scope. The 409-query training set and the 50-query frozen primary evaluation set",
            LIM_DATA,
        ),
        (
            "Roman Urdu coverage. The transliteration dictionary, despite a 6x expansion, remains a finite, manually curated lookup table.",
            LIM_ROMAN + " " + LIM_METRICS + " " + LIM_HUMAN,
        ),
    ]
    hit, miss = replace_prefixed(doc, jobs)
    n_ins = insert_blocks_before_chapter6(doc)
    doc.save(str(SRC))
    print("replaced", len(hit), "miss", len(miss))
    for m in miss:
        print("MISS", m)
    print("inserted_m0_blocks", n_ins)


if __name__ == "__main__":
    main()
