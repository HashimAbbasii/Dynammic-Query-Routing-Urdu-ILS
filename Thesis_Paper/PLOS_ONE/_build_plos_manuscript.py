# -*- coding: utf-8 -*-
"""Build a PLOS ONE-formatted research article (double-spaced, line numbers, Vancouver)."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

HERE = Path(__file__).resolve().parent
OUT = HERE / "Shazad_Aslam_PLOS_ONE_Urdu_query_routing.docx"
COVER = HERE / "Shazad_Aslam_PLOS_ONE_cover_letter.docx"
FIG_SRC = Path(r"c:\Users\User\OneDrive\Documents\ULTRA_Project\validate\dual_index_routing\figures")
FIG_DST = HERE / "figures"


def set_run_font(run, size=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_line_numbers(section):
    sectPr = section._sectPr
    el = OxmlElement("w:lnNumType")
    el.set(qn("w:countBy"), "1")
    el.set(qn("w:restart"), "continuous")
    sectPr.append(el)


def p(doc, text, *, style="Normal", size=12, bold=False, italic=False, center=False, space_after=0, first_line=True):
    para = doc.add_paragraph()
    if style and style in doc.styles:
        try:
            para.style = style
        except Exception:
            pass
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if first_line and not center and style == "Normal":
        pf.first_line_indent = Inches(0.5)
    else:
        pf.first_line_indent = Inches(0)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return para


def heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Inches(0)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.keep_with_next = True
    size = 14 if level == 1 else 12
    run = para.add_run(text)
    set_run_font(run, size=size, bold=True, italic=(level == 2 and False))
    if level == 2:
        set_run_font(run, size=12, bold=True, italic=True)
    if level == 3:
        set_run_font(run, size=12, bold=False, italic=True)
    return para


def caption(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Inches(0)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(text)
    set_run_font(run, size=11, italic=True)
    return para


def add_table(doc, headers, rows, col_w=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run_font(r, size=10, bold=True)
        cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            set_run_font(r, size=10)
            cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    return table


def build_cover():
    d = Document()
    sec = d.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1))
    p(d, "PLOS ONE Editorial Office", first_line=False)
    p(d, "Cover letter — research article", first_line=False, italic=True)
    p(d, "", first_line=False)
    p(
        d,
        "Please consider this research article, “Headline or article? Replacing a character-length switch with learned query routing for Urdu news search.”",
        first_line=False,
    )
    p(
        d,
        "The manuscript comes out of an MS thesis at Air University, Islamabad. It is not a claim that we beat a simple baseline on every metric. The main point is narrower. ULTRA still routes Urdu queries with a 150-character cutoff. We treat SHORT and LONG as retrieval need (headline enough vs need the body), send those decisions to two indexes, and freeze the test sets.",
        first_line=False,
    )
    p(
        d,
        "On 40 trap queries the SVM beats word count 60% to 20%, but only where why/how/fact wording is present. On graded P@5 the SVM does not win (33.00% vs 36.50% for word count). nDCG@5 is highest if every query stays in the headline index. We kept those numbers in the paper on purpose. PLOS ONE asks for technically sound work, not a guaranteed retrieval breakthrough, and this record is complete enough to replicate.",
        first_line=False,
    )
    p(d, "Article type: Research Article. No prior PLOS submissions. We have no opposed reviewers.", first_line=False)
    p(d, "Suggested editors: anyone handling information retrieval, multilingual NLP, or low-resource search.", first_line=False)
    p(d, "Sincerely,", first_line=False)
    p(d, "Hashim Shazad (corresponding author)", first_line=False)
    p(d, "Department of Creative Technologies, Air University, Islamabad, Pakistan", first_line=False)
    p(d, "Supervisor: Dr. Adnan Aslam", first_line=False)
    d.save(str(COVER))


def build_paper():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    add_line_numbers(sec)
    sec.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sec.footer.paragraphs[0].add_run("Page ")
    set_run_font(run, size=10)
    # PAGE field
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    sec.footer.paragraphs[0]._p.append(fld)

    # Title page
    p(doc, "Headline or article? Replacing a character-length switch with learned query routing for Urdu news search",
      size=16, bold=True, center=True, first_line=False)
    p(doc, "Short title: Learned routing for Urdu news search", italic=True, center=True, first_line=False)
    p(doc, "Hashim Shazad1*, Adnan Aslam1", center=True, first_line=False)
    p(doc, "1 Department of Creative Technologies, Air University, Islamabad, Pakistan", center=True, first_line=False)
    p(doc, "* Corresponding author: Hashim Shazad, Department of Creative Technologies, Air University, Islamabad, Pakistan. Email to be supplied at submission.",
      center=True, first_line=False)
    p(doc, "ORCID (corresponding author): to be added in the PLOS submission system.", italic=True, center=True, first_line=False)

    heading(doc, "Abstract", 1)
    p(doc,
      "Urdu search mixes native script with Roman Urdu typed on a Latin keyboard. A recent news retriever, ULTRA, still decides how to treat a query with one cutoff: 150 characters. That number is easy to code and easy to misunderstand. A short question can need a full article. A long question can be asking for a single fact that already sits in a headline.",
      first_line=False)
    p(doc,
      "We kept the SHORT/LONG split but changed what it means. SHORT means a headline is probably enough. LONG means the user likely needs the article. An eight-feature support vector machine makes the call. Headlines and full articles live in two indexes. Confidence acts like a traffic light: high confidence searches one index, medium mixes both, low expands the query and then mixes.",
      first_line=False)
    p(doc,
      "Three frozen tests are reported separately. On 50 held-out queries the SVM hit 86% label accuracy and a six-word rule hit 84% (McNemar p = 1.0). On 40 trap queries written so that length and need disagree, the SVM reached 60%, word count 20%, and the 150-character rule 50% (McNemar 16–0, p < 0.001). Split by wording, the SVM was 18/18 when why/how/fact cues fired and 6/22 (27.27%) when they did not — the same as word count on that slice. Four hundred graded judgments on the same 40 queries did not reward the SVM at P@5 (word count 36.50%, always-headline and θ = 150 both 35.00%, always-full 34.25%, SVM 33.00%). nDCG@5 was highest for always-headline (0.6868).",
      first_line=False)
    p(doc,
      "The character cutoff is a weak proxy for retrieval need. The SVM helps when the query carries an obvious cue. On this news collection it does not improve early precision. We report that result as part of the study, not as a footnote.",
      first_line=False)

    heading(doc, "Author summary", 1)
    p(doc,
      "If you search Urdu news, the system has to guess whether you want a headline or the story under it. Counting letters is a poor guess. We trained a small model to make that guess and we gave it two rooms to search. It is better than counting words when the query says why or how. It is not better at putting a useful article in the first five hits. That second fact is the one we would have liked to hide. We did not.",
      first_line=False)

    heading(doc, "Introduction", 1)
    p(doc,
      "Most English retrieval papers can assume that a query is typed in one script and that “short” and “long” roughly track how much context the user gave. Urdu search does not work like that. People switch between Perso-Arabic Urdu and Roman Urdu in the same session, spell the same word several ways, and pack a request for an explanation into three tokens [1–3].")
    p(doc,
      "ULTRA is a recent dual-embedding news search system built for that setting [4]. Its routing rule is a tape measure: queries with 150 characters or more take one path, the rest take another. We started from that system because it is a published, inspectable baseline, not because we think character length is a theory of information need.")
    p(doc,
      "A six-word cutoff is the other obvious tape. It is what a tired engineer would ship. It also fails on the queries that actually hurt: “team haar ka tajzia” is short and needs the article; “aaj pakistan ka score kya hai” can look long and still be a headline question.")
    p(doc,
      "Learned routers exist for English. Arabzadeh and colleagues trained a classifier to pick sparse, dense, or hybrid retrieval per query [5]. Jeong et al. route retrieval-augmented generation by predicted question complexity [6]. Both lines of work were built on English resources. Urdu now has better evaluation material than it did ten years ago — CURE, an Urdu MS MARCO translation — but those collections still do not tell a system which index to open [7,8].")
    p(doc,
      "This paper asks a smaller question than “can we beat GPT-4 at routing.” We ask whether SHORT versus LONG can be treated as headline-enough versus need-the-article, whether a cheap SVM can learn that on Urdu and Roman Urdu, and whether wiring the decision into two indexes actually moves P@5. The first two parts mostly work when the wording cooperates. The third, on the frozen traps we care about, does not.")

    heading(doc, "Materials and methods", 1)
    heading(doc, "Ethics and data", 2)
    p(doc,
      "No human-subjects protocol was required for news text already published on the web. Relevance labels were produced by the authors under a written rubric (headline enough vs need the article). They are not two independent paid annotators. We say that in the limitations because a journal reader will ask.")

    heading(doc, "Collection and indexes", 2)
    p(doc,
      "The corpus is the ULTRA news collection: 111,860 Urdu articles after cleaning, stored in ChromaDB collection urdu_news. That collection is the full-article room. The headline room is a separate dense index over titles, with embeddings cached for reuse. Both rooms use the same multilingual MiniLM encoder (paraphrase-multilingual-MiniLM-L12-v2) [9]. We did not retrain the encoder. If retrieval later fails, that is part of the result, not an excuse to swap models after seeing scores.")
    p(doc,
      "Roman Urdu queries are mapped through a word list of 198 pairs on disk (older drafts said 179; the file on disk is 198). Unlisted tokens stay as typed. Native-script Urdu skips this step.")

    heading(doc, "Router", 2)
    p(doc,
      "The frozen generalization model (Phase 3B V2) is an RBF-kernel SVM on eight features: Urdu-character ratio, Roman/Latin ratio, has-Urdu, has-Roman, token count, character count, mixed-script flag, and Urdu-character count. Features are z-scored with the training scaler. A later twelve-feature pickle exists in the deployed demo; we do not use it as the Phase 3B number, because that pickle saw extra trap-style training mass. Mixing those two models is how a 96% figure appears. We do not report 96%.")
    p(doc,
      "Confidence is the SVM probability of the chosen class. HIGH is 85% or above: search only the chosen room. MEDIUM is 60% up to 85%: mix both rooms at equal weight. LOW is below 60%: expand the query slightly, then mix. The expansion path is implemented. On the live demo queries we used for defense, LOW did not fire. We do not pretend that it did.")

    heading(doc, "What SHORT and LONG mean here", 2)
    p(doc,
      "Labels are not “is the string short.” SHORT means a competent headline could answer the user (score, price, date, who won, did the event happen). LONG means the user asked for why, how, impact, comparison, or a story. The 40 trap queries were written to that rule and then frozen. They were not used to train the V2 SVM.")

    heading(doc, "Evaluation layers", 2)
    p(doc,
      "We keep three layers apart, because they answer different questions.")
    p(doc,
      "Development and cross-validation show that the routing task is learnable. Some splits reach 100% against 50% for θ = 150. That is not a generalization claim. Anyone who has overfit a small Urdu query set knows how that 100% happens.")
    p(doc,
      "Frozen Phase 3B uses 50 primary queries held out from V2 training (409 labeled queries in that training pool: 193 SHORT, 216 LONG). Gold labels were fixed before scoring. The word-count baseline (≥ 6 tokens → LONG) was also frozen.")
    p(doc,
      "Frozen traps H001–H040 are 40 queries built so that length and need disagree, plus easy controls. Dual-index P@5 uses the same 40, depth 5, 400 graded judgments (0/1/2). A smaller earlier pilot (Phase 2.5, 33 queries) is reported only as supporting context.")
    p(doc,
      "Classification accuracy, exact McNemar tests, P@5, and nDCG@5 are the metrics. We do not convert a non-significant McNemar p-value into a story about “comparable superiority.”")

    heading(doc, "Results", 1)
    heading(doc, "Phase 3B: the SVM is barely ahead of counting words", 2)
    p(doc,
      "On the 50 primary queries the V2 SVM was correct on 43 (86%). Word count was correct on 42 (84%). The McNemar table is 2 vs 1 discordant pairs, exact p = 1.0. If this were the only experiment in the paper, we would not have much to say. It is useful mainly as a check that the frozen extractor matches the model that was actually trained.")

    heading(doc, "Trap queries: a real classification gap, with a catch", 2)
    p(doc,
      "On H001–H040 the SVM reached 60%, word count 20%, and θ = 150 50%. McNemar against word count is 16–0, p < 0.001. That is the result we would put on a slide if we were being careless.")
    p(doc,
      "The cue split is the part that should stay on the slide. Eighteen queries fire why/how/fact features. The SVM got 18/18; word count got 2/18. The other 22 queries do not fire those cues. Both systems sit at 27.27%. So the trap win is not “the SVM understands retrieval need.” It is “the SVM notices certain words, and the traps were rich in those words.”")

    add_table(
        doc,
        ["Layer", "n", "SVM", "Word count ≥ 6", "θ = 150"],
        [
            ["Development / CV (learnability only)", "varies", "100% on some splits", "—", "50%"],
            ["Frozen Phase 3B (V2, 8 features)", "50", "86% (43/50)", "84% (42/50)", "—"],
            ["Frozen traps H001–H040", "40", "60%", "20%", "50%"],
            ["Traps with why/how/fact cue", "18", "100% (18/18)", "11.11% (2/18)", "—"],
            ["Traps with no cue", "22", "27.27%", "27.27%", "—"],
        ],
    )
    caption(doc, "Table 1. Classification accuracy by evaluation layer. Development/CV is not a generalization result. Phase 3B McNemar p = 1.0. Trap McNemar vs word count: 16–0, p < 0.001.")

    heading(doc, "Retrieval: the classification win does not transfer", 2)
    p(doc,
      "Table 2 is the number that almost got edited out of an early draft. On the same 40 queries, with two indexes live, word count had the best P@5 (36.50%). Always-headline and θ = 150 tied at 35.00% (every query was SHORT under the character rule). Always-full-article was 34.25%. The SVM was last at 33.00%. nDCG@5 ranked always-headline first (0.6868), then word count (0.6476), SVM (0.6149), always-full (0.6020).")
    p(doc,
      "A smaller dual-index pilot on 33 judged queries (Phase 2.5) had a tiny SVM edge at P@5 (35.76% vs 35.15% word count vs 32.73% for θ = 150). Depth stopped at 5. We do not stack that pilot on top of the frozen 40 to invent a retrieval win.")

    add_table(
        doc,
        ["Router", "P@5", "nDCG@5", "Queries → headline", "Queries → full article"],
        [
            ["Word count ≥ 6", "36.50%", "0.6476", "20", "20"],
            ["Always headline / θ = 150", "35.00%", "0.6868", "40", "0"],
            ["Always full article", "34.25%", "0.6020", "0", "40"],
            ["SVM (V2)", "33.00%", "0.6149", "20", "20"],
        ],
    )
    caption(doc, "Table 2. Frozen dual-index retrieval on H001–H040 (400 graded judgments, cutoff 5). Higher is better. The SVM does not win P@5 or nDCG@5.")

    p(doc,
      "Fig 1 is the system we actually ran: dictionary, SVM, two rooms, lights. Fig 2 is the cue split. Fig 3 is the P@5 comparison. Fig 4 is a reminder not to quote development 100% as if it were Table 1.")
    caption(doc, "Fig 1. Routing pipeline. Roman Urdu dictionary, SVM room choice (headline vs full article), and HIGH / MEDIUM / LOW mixing.")
    caption(doc, "Fig 2. Trap cue split (n = 40). SVM 18/18 when why/how/fact wording fires; 27.27% for both systems otherwise.")
    caption(doc, "Fig 3. Dual-index graded P@5 on the frozen 40 queries.")
    caption(doc, "Fig 4. Three evaluation layers that should not be averaged into one headline number.")

    heading(doc, "Live confidence, briefly", 2)
    p(doc,
      "A defense demo on the deployed twelve-feature pickle (not the frozen V2 scorer) produced a green light on کرکٹ میچ (about 99%, headlines only) and yellow lights on ڈالر کی قیمت کتنی بڑھی (about 83%) and آج سٹاک ایکسچینج کتنے پوائنٹ پر (about 66%). A red light is in the code. It did not show up on that small probe. Anyone repeating the demo should not recycle an older JSON that put a cricket-score query in the red bin; on the current pickle that query is green.")

    heading(doc, "Discussion", 1)
    p(doc,
      "There are two honest readings of this project, and they are not the same paper.")
    p(doc,
      "Reading A: character length is the wrong switch for Urdu news queries. A small SVM, given script and cue features, can beat a word-count rule on queries that were written to punish that rule. We believe Reading A. The 16–0 McNemar table is hard to argue with, provided one also prints the 18/22 cue split.")
    p(doc,
      "Reading B: therefore retrieval gets better. We do not have Reading B. Once both rooms are real indexes, sending a query to the “right” room on our labels did not raise P@5. Always searching headlines even won nDCG@5. That can mean several things at once: the gold labels are about need, not about which room this encoder happens to rank well; news headlines already carry a lot of the lexical overlap MiniLM likes; mixing rooms on medium confidence can import junk; 40 queries is a small IR test. All of those can be true.")
    p(doc,
      "English adaptive retrieval work usually reports ranking metrics as the main figure and treats the router as a means [5,6]. If we had followed that habit and led with 86% vs 84%, or with development 100%, the paper would look cleaner and be worse. PLOS ONE’s bar is technical soundness [10]. A negative retrieval transfer on a frozen set is still a result.")
    p(doc,
      "The practical implication for Urdu news search is modest. Do not ship θ = 150 and call it adaptive. Do not ship our SVM and promise better first-page results. If the product need is “don’t embarrass yourself on why/how queries,” the cue features are doing real work. If the product need is nDCG, start with the headline index and spend the next month on the encoder and the judgments, not on another kernel.")

    heading(doc, "Limitations", 2)
    p(doc,
      "The corpus is news. Legal or medical Urdu would look different. Roman Urdu coverage is a word list, not a trained transliterator, so unseen spellings pass through untouched. Relevance labels were not a formal double-blind annotation study; agreement with a second independent rater is not reported because that rater does not exist. The deployed pickle and the frozen V2 model are not the same object; we documented that instead of averaging them. LOW-confidence expansion rarely triggers on the queries we tried. P@5 at depth 5 cannot see a relevant article sitting at rank 8.")

    heading(doc, "Conclusions", 1)
    p(doc,
      "We replaced ULTRA’s 150-character switch with a need-based SHORT/LONG label, an eight-feature SVM, two indexes, and a confidence mixer. The SVM is easy to beat word count with when the query contains why/how/fact wording, and it is not easy when those words are absent. Wiring the decision into retrieval did not improve P@5 on 40 frozen traps. nDCG@5 favored searching headlines for every query. That is the study. It is enough to stop quoting character thresholds as if they were a theory of the user, and not enough to declare the retrieval problem solved.")

    heading(doc, "Acknowledgments", 1)
    p(doc,
      "We thank Air University’s Department of Creative Technologies for the environment in which this MS work was done. Named editors and reviewers are not listed here, per PLOS policy.")

    heading(doc, "References", 1)
    refs = [
        "1. Daud A, Khan W, Che D. Urdu language processing: a survey. Artif Intell Rev. 2017;47: 279–311.",
        "2. Kazi S, Khoja S. A survey of Urdu NLP: resources, models, and challenges. ACM Comput Surv. 2025.",
        "3. Hussain S, et al. Roman Urdu processing for downstream NLP tasks. 2025.",
        "4. Bashir M, Qaiser S, Hussain S. ULTRA: dual-embedding query-length routing for Urdu news retrieval. 2026.",
        "5. Arabzadeh N, et al. Query-level routing between sparse and dense retrievers. Proceedings of SIGIR. 2021.",
        "6. Jeong S, et al. Adaptive-RAG: learning to adapt retrieval-augmented large language models. Proceedings of NAACL. 2024.",
        "7. Iqbal H, et al. CURE: Collection for Urdu Retrieval Evaluation. 2021.",
        "8. Butt M, et al. Urdu MS MARCO: a translated passage-ranking benchmark. 2024.",
        "9. Reimers N, Gurevych I. Sentence-BERT: sentence embeddings using Siamese BERT-networks. Proceedings of EMNLP. 2019.",
        "10. PLOS ONE. Publication criteria. Available from: https://journals.plos.org/plosone/s/criteria-for-publication",
        "11. Robertson S, Zaragoza H. The probabilistic relevance framework: BM25 and beyond. Found Trends Inf Retr. 2009;3(4): 333–389.",
        "12. Karpukhin V, Oguz B, Min S, Lewis P, Wu L, Edunov S, et al. Dense passage retrieval for open-domain question answering. Proceedings of EMNLP. 2020. pp. 6769–6781.",
        "13. Conneau A, Khandelwal K, Goyal N, Chaudhary V, Wenzek G, Guzmán F, et al. Unsupervised cross-lingual representation learning at scale. Proceedings of ACL. 2020. pp. 8440–8451.",
        "14. Platt J. Probabilistic outputs for support vector machines and comparisons to regularized likelihood methods. In: Advances in large margin classifiers. 1999. pp. 61–74.",
        "15. McNemar Q. Note on the sampling error of the difference between correlated proportions or percentages. Psychometrika. 1947;12: 153–157.",
    ]
    for r in refs:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Inches(-0.4)
        para.paragraph_format.left_indent = Inches(0.4)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(r)
        set_run_font(run, size=12)

    heading(doc, "Supporting information", 1)
    p(doc, "S1 File. Frozen trap sheet and label protocol (SHORT = headline enough, LONG = need the article).", first_line=False)
    p(doc, "S2 File. Held-out dual-index P@5 and nDCG@5 tables (400 judgments).", first_line=False)
    p(doc, "S3 File. Phase 3B frozen classification report (50 primary queries).", first_line=False)

    heading(doc, "Declarations (enter in the PLOS submission form as well)", 1)
    p(doc, "Data availability: Frozen scores and label sheets are in the project repository under validate/dual_index_routing/labels/. The news corpus and Chroma index are too large for the journal upload; they can be shared with qualified researchers on request to the corresponding author, subject to the original crawl’s terms.", first_line=False)
    p(doc, "Funding: The authors received no specific funding for this work. It was completed as part of an MS degree.", first_line=False)
    p(doc, "Competing interests: The authors have declared that no competing interests exist.", first_line=False)
    p(doc, "Author contributions: Hashim Shazad: conceptualization, software, investigation, data curation, writing – original draft. Adnan Aslam: supervision, writing – review and editing.", first_line=False)

    doc.save(str(OUT))
    print("wrote", OUT)


def copy_figures():
    FIG_DST.mkdir(parents=True, exist_ok=True)
    mapping = {
        "fig_two_rooms_lights.png": "Fig1_routing_pipeline.png",
        "fig_cue_split.png": "Fig2_cue_split.png",
        "fig_heldout_p5.png": "Fig3_p5.png",
        "fig_three_evaluation_layers.png": "Fig4_evaluation_layers.png",
    }
    for src_name, dst_name in mapping.items():
        src = FIG_SRC / src_name
        if src.exists():
            shutil.copy2(src, FIG_DST / dst_name)
            print("copied", dst_name)
        else:
            print("missing", src)


def main():
    FIG_DST.mkdir(parents=True, exist_ok=True)
    copy_figures()
    build_cover()
    build_paper()
    print("cover", COVER)


if __name__ == "__main__":
    main()
